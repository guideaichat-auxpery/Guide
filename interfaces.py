import streamlit as st
from utils import call_openai_api, get_max_tokens_for_user_type, scroll_to_top, add_scroll_to_top_button, render_conversation_sidebar, manage_conversation_history, apply_chatgpt_chat_style, scroll_chat_to_bottom, inject_chat_auto_scroll, validate_file_upload, sanitize_filename
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import io
import uuid
import json
import os
from datetime import datetime
from database import get_db, log_student_activity, database_available

def show_lesson_planning_interface():
    """Educational planning interface for educators with Australian Curriculum alignment"""
    apply_chatgpt_chat_style()
    inject_chat_auto_scroll()
    
    from utils import render_conversation_sidebar
    
    # Get educator info
    user_id = st.session_state.get('user_id')
    
    # Initialize session-specific conversation ID and conversation if not exists
    if 'planning_session_id' not in st.session_state:
        # Try to auto-load most recent conversation first
        if database_available and user_id:
            from database import get_user_chat_conversations, create_chat_conversation, load_conversation_to_session
            db = get_db()
            if db:
                try:
                    # Get existing conversations
                    existing_conversations = get_user_chat_conversations(
                        db, user_id=user_id, interface_type='planning'
                    )
                    
                    if existing_conversations and len(existing_conversations) > 0:
                        # Auto-load most recent conversation
                        most_recent = existing_conversations[0]
                        st.session_state.planning_session_id = most_recent.session_id
                        st.session_state['planning_current_conversation_id'] = most_recent.id
                        
                        # Load messages from database
                        loaded_messages = load_conversation_to_session(
                            db, most_recent.session_id, 'planning'
                        )
                        if loaded_messages:
                            st.session_state.planning_messages = loaded_messages
                            # Show restore notification
                            restore_time = most_recent.last_activity.strftime('%d/%m/%Y %H:%M') if hasattr(most_recent, 'last_activity') and most_recent.last_activity else 'earlier'
                            st.toast(f"✓ Restored your conversation from {restore_time}", icon="🔄")
                    else:
                        # Create first conversation for new users
                        st.session_state.planning_session_id = str(uuid.uuid4())
                        title = f"Planning {datetime.now().strftime('%d/%m %H:%M')}"
                        new_conv = create_chat_conversation(
                            db, title=title, session_id=st.session_state.planning_session_id,
                            interface_type='planning', user_id=user_id, student_id=None
                        )
                        st.session_state['planning_current_conversation_id'] = new_conv.id
                except Exception as e:
                    print(f"Error loading/creating planning conversation: {str(e)}")
                    st.session_state.planning_session_id = str(uuid.uuid4())
                finally:
                    db.close()
        else:
            st.session_state.planning_session_id = str(uuid.uuid4())
    
    # Render conversation sidebar (handles conversation management)
    if database_available and user_id:
        render_conversation_sidebar('planning', user_id=user_id)
    
    # Ensure planning messages are initialized separately
    if 'planning_messages' not in st.session_state:
        st.session_state.planning_messages = []
        
    st.markdown("### 📚 Montessori Educational Planning Tool")
    st.markdown("*Create comprehensive lesson plans with Montessori National Curriculum and Australian Curriculum V.9 alignment*")
    
    # Age group selector with year level display
    from utils import map_age_to_year_levels
    
    age_group = st.selectbox(
        "Select Age Group:",
        ["12-15", "9-12", "6-9", "3-6"],
        format_func=lambda x: {
            "3-6": "Early Years (3-6) → Foundation",
            "6-9": "Lower Primary (6-9) → Years 1-3", 
            "9-12": "Upper Primary (9-12) → Years 4-6",
            "12-15": "Adolescent (12-15) → Years 7-9"
        }[x]
    )
    
    # Year level mapping (not displayed to user)
    year_levels = map_age_to_year_levels(age_group)
    
    # Planning type selector
    planning_type = st.selectbox(
        "Planning Type:",
        ["lesson_plan", "assessment_rubric"],
        format_func=lambda x: {
            "lesson_plan": "Lesson Planning",
            "assessment_rubric": "Assessment Rubric"
        }[x]
    )

    # Mode toggle — Generate New Plan vs Align My Plan
    planning_mode = st.radio(
        "Mode:",
        ["✨ Generate New Plan", "📄 Align My Plan"],
        horizontal=True,
        key="planning_mode_toggle"
    )

    st.markdown("---")

    # ---- ALIGN MY PLAN MODE ----
    if planning_mode == "📄 Align My Plan":
        st.markdown("#### Upload Your Lesson Plan or Task Sheet")

        align_descriptions = {
            "3-6": (
                "*Upload a document you've already created and the AI will identify the Montessori connections "
                "and Australian Curriculum V9 (Foundation) alignment within it — affirming what you've done well, "
                "highlighting Sensitive Periods and Prepared Environment connections, and suggesting practical "
                "Montessori enhancements for the Early Years.*"
            ),
            "6-9": (
                "*Upload a document you've already created and the AI will identify the Montessori connections "
                "and Australian Curriculum V9 (Years 1–3) alignment within it — affirming what you've done well, "
                "highlighting Cosmic Education and Great Stories connections, and suggesting concrete-to-abstract "
                "Montessori enhancements for the Second Plane.*"
            ),
            "9-12": (
                "*Upload a document you've already created and the AI will identify the Montessori connections "
                "and Australian Curriculum V9 (Years 4–6) alignment within it — affirming what you've done well, "
                "highlighting opportunities for inquiry, Going Out, and real-world purpose, and suggesting "
                "Montessori enhancements suited to the upper Second Plane.*"
            ),
            "12-15": (
                "*Upload a document you've already created and the AI will identify the Montessori connections "
                "and Australian Curriculum V9 (Years 7–9) alignment within it — affirming what you've done well, "
                "suggesting Montessori enhancements, and mapping your plan to the four stages of the "
                "GUIDE Learning Design Protocol.*"
            ),
        }
        st.markdown(align_descriptions.get(age_group, align_descriptions["9-12"]))

        uploaded_plan = st.file_uploader(
            "Upload your plan (PDF or Word document):",
            type=["pdf", "docx"],
            key="align_plan_upload",
            help="Accepts PDF (.pdf) and Word documents (.docx)"
        )

        extra_context = st.text_area(
            "Any extra context? (optional)",
            placeholder="e.g. This is for a Year 5 class focused on sustainability. We have 3 sessions of 90 minutes.",
            key="align_plan_context",
            height=90
        )

        analyse_clicked = st.button("🔍 Analyse & Align", type="primary", use_container_width=True)

        if analyse_clicked:
            if not uploaded_plan:
                st.warning("Please upload a PDF or Word document first.")
            else:
                from utils import validate_file_upload, get_alignment_system_prompt, call_openai_api
                import PyPDF2
                import io as _io

                is_valid, error_msg = validate_file_upload(uploaded_plan)
                if not is_valid:
                    st.error(f"File validation failed: {error_msg}")
                else:
                    document_content = ""
                    with st.spinner("Reading your document..."):
                        if uploaded_plan.type == "application/pdf":
                            pdf_reader = PyPDF2.PdfReader(_io.BytesIO(uploaded_plan.read()))
                            extracted_pages = [page.extract_text() or "" for page in pdf_reader.pages]
                            document_content = "\n".join(extracted_pages)
                            if not document_content.strip():
                                st.error("This PDF appears to be image-based and the text could not be extracted. Please try a Word document or a text-based PDF.")
                                st.stop()
                        elif uploaded_plan.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            from docx import Document as DocxDocument
                            doc = DocxDocument(_io.BytesIO(uploaded_plan.read()))
                            document_content = "\n".join([para.text for para in doc.paragraphs])

                    MAX_DOC_CHARS = 8000
                    was_truncated = len(document_content) > MAX_DOC_CHARS
                    document_content = document_content[:MAX_DOC_CHARS]
                    if was_truncated:
                        st.info("Your document was quite long — the first 8,000 characters were analysed. For best results, consider uploading a shorter excerpt or summary.")

                    from utils import extract_year_level_from_query, extract_subject_from_query
                    search_text = document_content + " " + (extra_context or "")
                    detected_year_level = extract_year_level_from_query(search_text)
                    detected_subject = extract_subject_from_query(search_text)

                    st.session_state.align_plan_document = document_content
                    st.session_state.align_plan_extra_context = extra_context

                    detection_header_parts = []
                    if detected_year_level:
                        detection_header_parts.append(f"Detected Year Level: {detected_year_level}")
                    if detected_subject:
                        detection_header_parts.append(f"Detected Subject: {detected_subject}")
                    detection_header = (", ".join(detection_header_parts) + "\n\n") if detection_header_parts else ""

                    user_message_parts = [
                        f"{detection_header}I'd like you to analyse and align the following lesson plan or task sheet to the Montessori curriculum and Australian Curriculum V9.\n\n**File:** {uploaded_plan.name}\n\n**Document Content:**\n{document_content}"
                    ]
                    if extra_context and extra_context.strip():
                        user_message_parts.append(f"\n\n**Additional Context from Educator:**\n{extra_context.strip()}")
                    user_message = "".join(user_message_parts)

                    system_prompt = get_alignment_system_prompt(age_group)

                    with st.spinner("Reading your plan and finding the curriculum connections..."):
                        messages = [{"role": "user", "content": user_message}]
                        result = call_openai_api(
                            messages=messages,
                            max_tokens=8000,
                            system_prompt=system_prompt,
                            is_student=False,
                            age_group=age_group,
                            interface_type="align_plan",
                            curriculum_type="Blended",
                            year_level=detected_year_level,
                            subject=detected_subject,
                            use_conversation_history=False
                        )

                    if result:
                        st.session_state.align_plan_result = result
                        st.session_state.align_plan_filename = uploaded_plan.name
                    else:
                        st.error("Something went wrong. Please try again.")

        # Display stored result
        if st.session_state.get("align_plan_result"):
            result_text = st.session_state.align_plan_result
            source_name = st.session_state.get("align_plan_filename", "your plan")

            st.markdown("---")
            st.markdown(f"#### Alignment Analysis — *{source_name}*")

            required_sections = [
                "What Your Plan Already Does Well",
                "Australian Curriculum V9 Alignment",
                "Montessori Connections"
            ]
            missing_sections = [s for s in required_sections if s.lower() not in result_text.lower()]
            if missing_sections:
                st.warning(
                    "The analysis may be incomplete — some sections could not be generated. "
                    "The response is shown below. You can try again if a section is missing."
                )
                if st.button("🔄 Try Again", key="align_retry_btn"):
                    stored_doc = st.session_state.get("align_plan_document", "")
                    stored_context = st.session_state.get("align_plan_extra_context", "")
                    if stored_doc:
                        from utils import get_alignment_system_prompt, call_openai_api, extract_year_level_from_query, extract_subject_from_query
                        search_text = stored_doc + " " + (stored_context or "")
                        retry_year_level = extract_year_level_from_query(search_text)
                        retry_subject = extract_subject_from_query(search_text)
                        retry_msg = f"I'd like you to analyse and align the following lesson plan or task sheet to the Montessori curriculum and Australian Curriculum V9.\n\n**Document Content:**\n{stored_doc}"
                        if stored_context:
                            retry_msg += f"\n\n**Additional Context from Educator:**\n{stored_context}"
                        with st.spinner("Trying again..."):
                            retry_result = call_openai_api(
                                messages=[{"role": "user", "content": retry_msg}],
                                max_tokens=8000,
                                system_prompt=get_alignment_system_prompt(age_group),
                                is_student=False,
                                age_group=age_group,
                                interface_type="align_plan",
                                curriculum_type="Blended",
                                year_level=retry_year_level,
                                subject=retry_subject,
                                use_conversation_history=False
                            )
                        if retry_result:
                            st.session_state.align_plan_result = retry_result
                            st.rerun()
                    else:
                        st.info("Please re-upload your document to try again.")

            st.markdown(result_text)

            st.markdown("---")
            st.markdown("**Export this alignment:**")
            col_pdf, col_docx = st.columns(2)

            from utils import export_lesson_plan_to_pdf, export_lesson_plan_to_docx

            with col_pdf:
                pdf_bytes = export_lesson_plan_to_pdf(
                    result_text,
                    title=f"Curriculum Alignment — {source_name}",
                    filename="curriculum_alignment.pdf"
                )
                if pdf_bytes:
                    st.download_button(
                        "⬇ Download PDF",
                        data=pdf_bytes,
                        file_name="curriculum_alignment.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

            with col_docx:
                docx_bytes = export_lesson_plan_to_docx(
                    result_text,
                    title=f"Curriculum Alignment — {source_name}",
                    filename="curriculum_alignment.docx"
                )
                if docx_bytes:
                    st.download_button(
                        "⬇ Download Word Doc",
                        data=docx_bytes,
                        file_name="curriculum_alignment.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            if st.button("Clear & Start Again", key="align_clear_btn"):
                st.session_state.pop("align_plan_result", None)
                st.session_state.pop("align_plan_filename", None)
                st.rerun()

        return

    # ---- GENERATE NEW PLAN MODE (unchanged) ----

    # Display chat history
    ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
    for message in st.session_state.planning_messages:
        avatar = ai_avatar if message["role"] == "assistant" else None
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
    
    # Scroll to bottom of chat after displaying messages
    if st.session_state.planning_messages:
        scroll_chat_to_bottom()
    
    # Check for pending follow-up prompt from button clicks
    pending_followup = st.session_state.pop('planning_followup_prompt', None)
    
    # Chat input for planning questions
    prompt = st.chat_input("Tell me your year level/age, topic, time, and instructions")
    
    # Use follow-up prompt if set, otherwise use chat input
    if pending_followup:
        prompt = pending_followup
    
    if prompt:
        st.session_state.planning_messages.append({
            "role": "user",
            "content": prompt
        })
        
        # Save user message to database
        save_success = False
        if database_available and user_id:
            from database import save_conversation_message
            from utils import update_conversation_title_if_needed
            db = get_db()
            if db:
                try:
                    save_conversation_message(
                        db,
                        session_id=st.session_state.planning_session_id,
                        interface_type='planning',
                        role='user',
                        content=prompt,
                        user_id=user_id,
                        student_id=None
                    )
                    save_success = True
                    
                    # Auto-title on first message
                    conv_id = st.session_state.get('planning_current_conversation_id')
                    if conv_id and len(st.session_state.planning_messages) == 1:
                        update_conversation_title_if_needed(db, conv_id, 'planning', prompt)
                except Exception as e:
                    print(f"Error saving planning user message: {str(e)}")
                    st.warning("⚠️ Unable to save message. Please check your connection.")
                finally:
                    db.close()
        
        # Show save confirmation if successful
        if save_success:
            st.toast("✓ Message saved", icon="💾")
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Get AI response
        ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
        with st.chat_message("assistant", avatar=ai_avatar):
            with st.spinner("Collaborating with you..."):
                # Construct system prompt based on planning type and age group
                if planning_type == "lesson_plan":
                    # Age-specific Montessori system context
                    montessori_context = {
                        "3-6": """IMPORTANT: Always use British English spelling and conventions (colour, organisation, analyse, centre, programme, etc.) in all responses.

You are creating a lesson plan for the First Plane of Development (ages 3-6 / Foundation / Cycle 1) following Montessori curriculum, pedagogy, and philosophy.

🎯 CRITICAL CURRICULUM RULE FOR AGES 3-6:
- PRIMARY FRAMEWORK: Montessori curriculum and pedagogy
- SECONDARY ALIGNMENT: Australian Curriculum (Foundation) may be referenced for alignment purposes ONLY
- When AC alignment is included, show the AC content descriptor codes (e.g., AC9EFLY01)
- The lesson design, materials, and approach MUST be based on Montessori principles, NOT Australian Curriculum

KEY MONTESSORI PRINCIPLES FOR AGES 3-6:
- Absorbent Mind: Children effortlessly absorb knowledge from their environment
- Sensitive Periods: Focus on order, language, movement, refinement of senses, and small objects
- Prepared Environment: Carefully organized materials that invite independence
- Practical Life exercises: Care of self, care of environment, grace and courtesy, control of movement
- Sensorial Materials: Pink Tower, Brown Stairs, Red Rods, Colour Tablets, Sound Cylinders, etc.
- Language Development: Sandpaper Letters, Movable Alphabet, Metal Insets
- Mathematics Foundation: Number Rods, Sandpaper Numbers, Spindle Boxes, Golden Beads
- Self-directed activity and freedom within limits
- Hands-on concrete materials before abstract concepts""",
                        
                        "6-9": """IMPORTANT: Always use British English spelling and conventions (colour, organisation, analyse, centre, programme, etc.) in all responses.

You are creating a lesson plan for the Second Plane of Development (ages 6-9 / Years 1-3 / Cycle 2) following Montessori curriculum, pedagogy, and philosophy.

🎯 CRITICAL CURRICULUM RULE FOR AGES 6-9:
- PRIMARY FRAMEWORK: Montessori curriculum and pedagogy
- SECONDARY ALIGNMENT: Australian Curriculum (Years 1-3) may be referenced for alignment purposes ONLY
- When AC alignment is included, show the AC content descriptor codes (e.g., AC9E3LA01)
- The lesson design, materials, and approach MUST be based on Montessori principles, NOT Australian Curriculum

KEY MONTESSORI PRINCIPLES FOR AGES 6-9:
- Reasoning Mind: Children develop abstract thinking and imagination
- Cosmic Education: Understanding the interconnectedness of all things and the child's role in the universe
- Great Stories: The Five Great Lessons as foundation for integrated learning
- Social Development: Group work, collaboration, peer learning
- Moral Development: Sense of justice, fairness, and societal rules
- Follow the Child: Child-led exploration within a structured framework
- Passage to Abstraction: From concrete materials to abstract reasoning
- Key Lessons: Story of the Universe, Coming of Life, Story of Humans, Story of Language, Story of Numbers
- Research and exploration become central to learning""",
                        
                        "9-12": """IMPORTANT: Always use British English spelling and conventions (colour, organisation, analyse, centre, programme, etc.) in all responses.

You are creating a lesson plan for the Second Plane of Development (ages 9-12 / Years 4-6 / Cycle 3) following Montessori curriculum, pedagogy, and philosophy.

🎯 CRITICAL CURRICULUM RULE FOR AGES 9-12:
- PRIMARY FRAMEWORK: Montessori curriculum and pedagogy
- SECONDARY ALIGNMENT: Australian Curriculum (Years 4-6) may be referenced for alignment purposes ONLY
- When AC alignment is included, show the AC content descriptor codes (e.g., AC9S5U01)
- The lesson design, materials, and approach MUST be based on Montessori principles, NOT Australian Curriculum

KEY MONTESSORI PRINCIPLES FOR AGES 9-12:
- Intellectual Independence: Deep research, critical thinking, and academic specialisation
- Cosmic Education Deepens: Understanding systems, interdependence, and contribution to society
- Hero Worship: Study of great contributors to civilisation
- Justice and Morality: Advanced ethical reasoning and social responsibility
- Cultural Subjects Integration: History, geography, biology, and physics interconnected
- Going Out: Extended field work and community engagement
- Follow the Child's Interests: Student agency in research topics and projects
- Timeline Work: Great lessons extended into detailed timeline exploration
- Entrepreneurial Spirit: Micro-economy, business ventures, real-world problem solving""",
                        
                        "12-15": """IMPORTANT: Always use British English spelling and conventions (colour, organisation, analyse, centre, programme, etc.) in all responses.

You are an expert curriculum designer for adolescent learners following the GUIDE Learning Design Protocol.
Your job is to design learning experiences that are human, inquiry-driven, and experiential — not traditional worksheet-based instruction.

🎯 THE GUIDE LEARNING DESIGN PROTOCOL
All learning designs must follow these four stages in order:

**STAGE 1 — ANCHOR IN MEANING**
Every unit must clearly define:
1. **Human Theme**: A universal concept relevant to being human (e.g., Identity, Power, Change, Belonging, Systems, Progress, Conflict, Truth)
2. **Driving Question**: A debatable, real-world question that frames the entire unit and creates intellectual tension
3. **Purpose Statement**: A one-sentence explanation of why this learning matters in the real world
If any of these three are missing, the design is incomplete.

**STAGE 2 — BUILD THE LEARNING STORY**
The unit must follow this five-phase narrative structure:
- **Hook & Tension**: Create curiosity, emotional engagement, or cognitive dissonance
- **Investigation**: Explicit teaching, research, experimentation, modelling, and skill development — always in service of the Driving Question
- **Construction**: Students create a meaningful product, argument, design, solution, or performance that responds to the Driving Question
- **Public Thinking**: Students share, defend, critique, revise, and articulate their thinking
- **Reflection**: Students explicitly describe what they now understand, how their thinking changed, and what they can now do

**STAGE 3 — DAILY SESSION DESIGN RULE**
Each learning session must answer:
- What thinking is being advanced today?
- Learning invitation (what learners are invited to do)
- What support or instruction is required?
- How will progress be visible today?
Sessions must always move the learning story forward.

**STAGE 4 — SYSTEM ALIGNMENT (LAST STEP ONLY)**
- Australian Curriculum V9 outcomes (Years 7-9) with AC content descriptor codes (e.g., AC9HH7K01, AC9E7LA01)
- Assessment language and reporting structures
- Learning design ALWAYS comes first, curriculum alignment is applied only after the experience is fully designed

🚫 NON-NEGOTIABLE CONSTRAINTS:
- Do NOT start from curriculum outcomes
- Do NOT design worksheet-driven lessons
- Do NOT generate activities without first establishing meaning and narrative flow
- ALWAYS prioritise inquiry, construction, and reflection

📋 OUTPUT FORMAT (Follow this order):
1. Human Theme
2. Driving Question
3. Purpose Statement
4. Learning Story (aligned with timeframe specified)
5. Session-by-Session Plan
6. Curriculum & Assessment Alignment (AC V9 codes)

✅ QUALITY CHECK (Before responding, internally verify):
Does this design feel like a real learning journey, not a school template? If not, revise.

MONTESSORI ADOLESCENT PRINCIPLES TO INTEGRATE:
- Erdkinder (Children of the Earth): Learning through meaningful work and community contribution
- Valorisation: Development of self-worth through authentic purpose
- Social Development: Peer relationships, identity formation, place in society
- Cosmic Education: Understanding systems, interdependence, and contribution to the world"""
                    }
                    system_context = montessori_context.get(age_group, "You are creating a lesson plan with Montessori principles and Australian Curriculum V9 alignment.")
                else:  # assessment_rubric
                    system_context = """You are creating an assessment rubric that balances Montessori observational assessment with Australian Curriculum V9 achievement standards.

IMPORTANT RUBRIC FORMAT REQUIREMENTS:
- Do NOT use letter grades (A, B, C, D) or numerical scores
- Use ONLY these four performance levels as column headers (left to right):
  1. Sophisticated
  2. High Expectation Met
  3. Expectation Met
  4. Developing
- For each criterion, describe what each performance level looks like
- Focus on descriptive, asset-based language that honors the child's development
- Align descriptions with AC V9 achievement standards where appropriate
- Include Montessori observational assessment approaches"""
                
                response = call_openai_api(
                    st.session_state.planning_messages,
                    system_prompt=system_context,
                    age_group=age_group,
                    year_level=year_levels[0] if year_levels else None,
                    curriculum_type="AC_V9",
                    interface_type="lesson_planning",
                    planning_type=planning_type
                )
                
                st.markdown(response)
                st.session_state.planning_messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                # Save assistant message to database
                assistant_save_success = False
                if database_available and user_id:
                    from database import save_conversation_message
                    db = get_db()
                    if db:
                        try:
                            save_conversation_message(
                                db,
                                session_id=st.session_state.planning_session_id,
                                interface_type='planning',
                                role='assistant',
                                content=response,
                                user_id=user_id,
                                student_id=None
                            )
                            assistant_save_success = True
                        except Exception as e:
                            print(f"Error saving planning assistant message: {str(e)}")
                            st.warning("⚠️ Unable to save response. Please check your connection.")
                        finally:
                            db.close()
                
                # Show save confirmation if successful
                if assistant_save_success:
                    st.toast("✓ Response saved", icon="💾")
                
                # Scroll to beginning of new response
                scroll_chat_to_bottom()
    
    # Export options
    if st.session_state.planning_messages:
        st.markdown("---")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Export as PDF", use_container_width=True):
                from utils import export_lesson_plan_to_pdf
                # Format conversation into content string
                content = ""
                for msg in st.session_state.planning_messages:
                    if msg['role'] == 'user':
                        content += f"**Educator Question:**\n{msg['content']}\n\n"
                    else:
                        content += f"{msg['content']}\n\n"
                
                # Create title from context
                title = f"{planning_type} - {age_group}"
                filename = f"lesson_plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                
                pdf_data, filename = export_lesson_plan_to_pdf(content, title, filename)
                st.download_button(
                    label="Download PDF",
                    data=pdf_data,
                    file_name=filename,
                    mime="application/pdf"
                )
        
        with col2:
            if st.button("📝 Export as DOCX", use_container_width=True):
                from utils import export_lesson_plan_to_docx
                # Format conversation into content string
                content = ""
                for msg in st.session_state.planning_messages:
                    if msg['role'] == 'user':
                        content += f"**Educator Question:**\n{msg['content']}\n\n"
                    else:
                        content += f"{msg['content']}\n\n"
                
                # Create title from context
                title = f"{planning_type} - {age_group}"
                filename = f"lesson_plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                docx_data, filename = export_lesson_plan_to_docx(content, title, filename)
                st.download_button(
                    label="Download DOCX",
                    data=docx_data,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Follow-up prompt suggestions for lesson plans
        if planning_type == "lesson_plan" and len(st.session_state.planning_messages) >= 2:
            st.markdown("---")
            st.markdown("**Continue your planning:**")
            follow_up_cols = st.columns(3)
            
            with follow_up_cols[0]:
                if st.button("📊 Design a Rubric", key="followup_rubric", use_container_width=True):
                    st.session_state.planning_followup_prompt = "Based on the lesson plan above, design a comprehensive assessment rubric using the four performance levels: Sophisticated, High Expectation Met, Expectation Met, and Developing. Include criteria that assess both the learning process and final product."
                    st.rerun()
            
            with follow_up_cols[1]:
                if st.button("📝 Student Templates", key="followup_templates", use_container_width=True):
                    st.session_state.planning_followup_prompt = "Create student work templates and scaffolds for this unit. Include templates for the Construction phase that help students organise their thinking, track their progress, and present their final product. Design these to support student agency while providing structure."
                    st.rerun()
            
            with follow_up_cols[2]:
                if st.button("🎯 Differentiation", key="followup_differentiation", use_container_width=True):
                    st.session_state.planning_followup_prompt = "Suggest differentiation strategies for this unit to support diverse learners. Include modifications for students who need additional support, extension opportunities for advanced learners, and strategies for different learning styles and needs."
                    st.rerun()
    
    # Add scroll to top button
    add_scroll_to_top_button()


def show_companion_interface():
    """Enhanced Montessori companion interface with conversation history management and persistence"""
    apply_chatgpt_chat_style()
    inject_chat_auto_scroll()
    
    from utils import manage_conversation_history, estimate_tokens, render_conversation_sidebar
    from database import save_conversation_message, log_educator_prompt, load_conversation_to_session
    
    # Get educator info
    user_id = st.session_state.get('user_id')
    
    # Initialize session-specific conversation ID and conversation if not exists
    if 'companion_session_id' not in st.session_state:
        # Try to auto-load most recent conversation first
        if database_available and user_id:
            from database import get_user_chat_conversations, create_chat_conversation
            db = get_db()
            if db:
                try:
                    # Get existing conversations
                    existing_conversations = get_user_chat_conversations(
                        db, user_id=user_id, interface_type='companion'
                    )
                    
                    if existing_conversations and len(existing_conversations) > 0:
                        # Auto-load most recent conversation
                        most_recent = existing_conversations[0]
                        st.session_state.companion_session_id = most_recent.session_id
                        st.session_state['companion_current_conversation_id'] = most_recent.id
                        
                        # Load messages from database
                        loaded_messages = load_conversation_to_session(
                            db, most_recent.session_id, 'companion'
                        )
                        if loaded_messages:
                            st.session_state.companion_messages = loaded_messages
                            # Show restore notification
                            restore_time = most_recent.last_activity.strftime('%d/%m/%Y %H:%M') if hasattr(most_recent, 'last_activity') and most_recent.last_activity else 'earlier'
                            st.toast(f"✓ Restored your conversation from {restore_time}", icon="🔄")
                    else:
                        # Create first conversation for new users
                        st.session_state.companion_session_id = str(uuid.uuid4())
                        title = f"Companion {datetime.now().strftime('%d/%m %H:%M')}"
                        new_conv = create_chat_conversation(
                            db, title=title, session_id=st.session_state.companion_session_id,
                            interface_type='companion', user_id=user_id, student_id=None
                        )
                        st.session_state['companion_current_conversation_id'] = new_conv.id
                except Exception as e:
                    print(f"Error loading/creating companion conversation: {str(e)}")
                    st.session_state.companion_session_id = str(uuid.uuid4())
                finally:
                    db.close()
        else:
            st.session_state.companion_session_id = str(uuid.uuid4())
    
    # Render conversation sidebar (handles conversation management)
    if database_available and user_id:
        render_conversation_sidebar('companion', user_id=user_id)
    
    # Ensure companion messages are initialized (will be populated by auto-load above if available)
    if 'companion_messages' not in st.session_state:
        st.session_state.companion_messages = []
    
    st.markdown("### 🗨️ Montessori Companion")
    st.markdown("*Your philosophical guide to Montessori principles, cosmic education, and educational wisdom*")
    
    # Age group selector for companion (optional - defaults to all ages)
    st.markdown("### 🌱 Select Age Group (Optional)")
    st.caption("Choose a specific age range for targeted guidance, or select 'All Ages' for comprehensive support")
    
    companion_age_options = {
        "All Ages (3-18)": "all",
        "Ages 12-15 (Adolescent)": "12-15",
        "Ages 9-12 (Upper Primary)": "9-12",
        "Ages 6-9 (Lower Primary)": "6-9",
        "Ages 3-6 (Early Years)": "3-6"
    }
    
    selected_age_display = st.selectbox(
        "Age Focus",
        options=list(companion_age_options.keys()),
        key="companion_age_selector",
        label_visibility="collapsed"
    )
    
    companion_age_group = companion_age_options[selected_age_display]
    
    # Manage conversation history (keep last 10 exchanges)
    st.session_state.companion_messages = manage_conversation_history(
        st.session_state.companion_messages, max_history=20
    )
    
    # Quick conversation starters - Comprehensive Montessori Guide Topics
    # MOVED TO TOP: Cards stay fixed here, chat flows below
    st.markdown("### 📚 Montessori Quick Guides")
    st.caption("Click any topic to explore authentic Montessori wisdom from Dr. Montessori's foundational texts")
    
    quick_prompts = [
        "🌌 What is cosmic education and how do I implement it?",
        "🎯 Explain the sensitive periods in child development",
        "🏛️ What is the prepared environment?",
        "👁️ How do I observe children effectively?",
        "🌱 What is the absorbent mind?",
        "✋ How do I introduce a new material?",
        "🔄 What are the three-period lessons?",
        "🌍 How does Montessori connect to the universe story?",
        "🔬 How do I implement cosmic education in science?",
        "📖 What are the great stories and how do I tell them?",
        "🤝 How do I handle social conflicts using Montessori principles?",
        "🌟 What is normalization and how do I recognize it?",
        "🧘 How do I create a culture of peace in the classroom?",
        "🎨 How does art connect to cosmic education?",
        "📊 How do I assess learning in a Montessori way?"
    ]
    
    cols = st.columns(3)
    for idx, prompt_text in enumerate(quick_prompts):
        with cols[idx % 3]:
            if st.button(prompt_text, key=f"quick_{idx}", use_container_width=True):
                # Add prompt to conversation
                st.session_state.companion_messages.append({
                    "role": "user",
                    "content": prompt_text
                })
                
                # Save to database if available
                if database_available and user_id:
                    db = get_db()
                    if db:
                        try:
                            save_conversation_message(
                                db,
                                session_id=st.session_state.companion_session_id,
                                interface_type='companion',
                                role='user',
                                content=prompt_text,
                                user_id=user_id,
                                student_id=None
                            )
                        except Exception as e:
                            print(f"Error saving conversation: {str(e)}")
                        finally:
                            db.close()
                
                st.rerun()
    
    st.markdown("---")
    
    # CRITICAL UI PLACEMENT: Document upload section MUST remain HERE (before chat history loop)
    # to prevent mid-conversation insertion during Streamlit reruns.
    # DO NOT move below chat history - it will appear between messages during multi-turn conversations.
    # Correct flow: Quick Guides → Upload Section → Chat History → Chat Input
    st.markdown("### 📁 Upload Teaching Materials for Feedback (Optional)")
    st.caption("Share lesson plans, observations, student work samples, or teaching materials for Montessori-focused feedback")
    
    uploaded_document = st.file_uploader(
        "Upload document for analysis",
        type=['txt', 'pdf', 'jpg', 'png', 'docx'],
        help="Upload teaching materials, lesson plans, observations, or student work for feedback",
        key="companion_document_upload"
    )
    
    # Process uploaded document and add to conversation
    if uploaded_document:
        # Validate file upload for security
        is_valid, error_msg = validate_file_upload(uploaded_document)
        if not is_valid:
            st.error(f"File validation failed: {error_msg}")
        elif st.button("🌟 Get Montessori Feedback", use_container_width=True, type="primary"):
            # Process document file
            document_content = ""
            with st.spinner("Reading your document..."):
                if uploaded_document.type == "application/pdf":
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_document.read()))
                    # Handle None values from scanned/textless pages
                    extracted_pages = [page.extract_text() or "" for page in pdf_reader.pages]
                    document_content = "\n".join(extracted_pages)
                    if not document_content.strip():
                        document_content = f"[PDF uploaded: {uploaded_document.name} - appears to be scanned/image-based. Text extraction not possible. Please describe the content or upload a text-based version.]"
                elif uploaded_document.type in ["image/jpeg", "image/png"]:
                    try:
                        image = Image.open(uploaded_document)
                        document_content = pytesseract.image_to_string(image)
                        if not document_content.strip():
                            document_content = f"[Image uploaded: {uploaded_document.name} - visual content not extractable as text]"
                    except:
                        document_content = f"[Image uploaded: {uploaded_document.name} - could not process]"
                elif uploaded_document.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = Document(io.BytesIO(uploaded_document.read()))
                    document_content = "\n".join([para.text for para in doc.paragraphs])
                else:
                    document_content = uploaded_document.read().decode("utf-8")
            
            # Create message with document content
            user_message = f"I've uploaded a document for your feedback: {uploaded_document.name}\n\n**Document Content:**\n{document_content}\n\nPlease provide Montessori-focused feedback on this material."
            
            st.session_state.companion_messages.append({
                "role": "user",
                "content": user_message
            })
            
            # Save to database
            if database_available and user_id:
                db = get_db()
                if db:
                    try:
                        save_conversation_message(
                            db,
                            session_id=st.session_state.companion_session_id,
                            interface_type='companion',
                            role='user',
                            content=user_message,
                            user_id=user_id,
                            student_id=None
                        )
                        st.toast("✓ Document uploaded and saved", icon="📄")
                    except Exception as e:
                        print(f"Error saving document message: {str(e)}")
                        st.warning("⚠️ Unable to save document. Please check your connection.")
                    finally:
                        db.close()
            
            st.rerun()
    
    st.markdown("---")
    
    # Check if last message needs a response (from Quick Guide click or chat input)
    need_response = (
        len(st.session_state.companion_messages) > 0 and 
        st.session_state.companion_messages[-1]["role"] == "user" and
        (len(st.session_state.companion_messages) == 1 or 
         st.session_state.companion_messages[-2]["role"] == "assistant")
    )
    
    # Display chat history
    ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
    for message in st.session_state.companion_messages:
        avatar = ai_avatar if message["role"] == "assistant" else None
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
    
    # Scroll to bottom of chat after displaying messages
    if st.session_state.companion_messages:
        scroll_chat_to_bottom()
    
    # If last message was from user (Quick Guide), generate response
    if need_response:
        ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
        with st.chat_message("assistant", avatar=ai_avatar):
            with st.spinner("Consulting Montessori wisdom..."):
                response = call_openai_api(
                    st.session_state.companion_messages,
                    age_group=companion_age_group if companion_age_group != "all" else None,
                    interface_type="companion"
                )
                st.markdown(response)
                st.session_state.companion_messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                # Scroll to beginning of new response
                scroll_chat_to_bottom()
                
                # Save assistant response to database
                assistant_save_success = False
                if database_available and user_id:
                    db = get_db()
                    if db:
                        try:
                            save_conversation_message(
                                db,
                                session_id=st.session_state.companion_session_id,
                                interface_type='companion',
                                role='assistant',
                                content=response,
                                user_id=user_id,
                                student_id=None
                            )
                            assistant_save_success = True
                        except Exception as e:
                            print(f"Error saving conversation: {str(e)}")
                            st.warning("⚠️ Unable to save response. Please check your connection.")
                        finally:
                            db.close()
                
                # Show save confirmation if successful
                if assistant_save_success:
                    st.toast("✓ Response saved", icon="💾")
    
    # Chat input
    if prompt := st.chat_input("Ask your Montessori question..."):
        st.session_state.companion_messages.append({
            "role": "user",
            "content": prompt
        })
        
        # Save user message to database
        save_success = False
        if database_available and user_id:
            from utils import update_conversation_title_if_needed
            db = get_db()
            if db:
                try:
                    save_conversation_message(
                        db,
                        session_id=st.session_state.companion_session_id,
                        interface_type='companion',
                        role='user',
                        content=prompt,
                        user_id=user_id,
                        student_id=None
                    )
                    
                    # Log for analytics
                    log_educator_prompt(
                        db, user_id, 'companion', prompt,
                        tokens_used=estimate_tokens(prompt)
                    )
                    save_success = True
                    
                    # Auto-title on first message
                    conv_id = st.session_state.get('companion_current_conversation_id')
                    if conv_id and len(st.session_state.companion_messages) == 1:
                        update_conversation_title_if_needed(db, conv_id, 'companion', prompt)
                except Exception as e:
                    print(f"Error saving conversation: {str(e)}")
                    st.warning("⚠️ Unable to save message. Please check your connection.")
                finally:
                    db.close()
        
        # Show save confirmation if successful
        if save_success:
            st.toast("✓ Message saved", icon="💾")
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Get AI response
        ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
        with st.chat_message("assistant", avatar=ai_avatar):
            with st.spinner("Consulting Montessori wisdom..."):
                response = call_openai_api(
                    st.session_state.companion_messages,
                    age_group=companion_age_group if companion_age_group != "all" else None,
                    interface_type="companion"
                )
                st.markdown(response)
                st.session_state.companion_messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                # Scroll to beginning of new response
                scroll_chat_to_bottom()
                
                # Save assistant response to database
                assistant_save_success = False
                if database_available and user_id:
                    db = get_db()
                    if db:
                        try:
                            save_conversation_message(
                                db,
                                session_id=st.session_state.companion_session_id,
                                interface_type='companion',
                                role='assistant',
                                content=response,
                                user_id=user_id,
                                student_id=None
                            )
                            assistant_save_success = True
                        except Exception as e:
                            print(f"Error saving conversation: {str(e)}")
                            st.warning("⚠️ Unable to save response. Please check your connection.")
                        finally:
                            db.close()
                
                # Show save confirmation if successful
                if assistant_save_success:
                    st.toast("✓ Response saved", icon="💾")
    
    # Add scroll to top button
    add_scroll_to_top_button()


def show_student_interface():
    """Enhanced student learning interface with curriculum context, conversation history, and persistence"""
    apply_chatgpt_chat_style()
    inject_chat_auto_scroll()
    
    from utils import manage_conversation_history, render_conversation_sidebar
    from database import save_conversation_message, load_conversation_to_session
    
    # Get student info from session
    student_id = st.session_state.get('user_id')
    student_name = st.session_state.get('user_name', 'Student')
    age_group = st.session_state.get('age_group', 'unknown')
    
    # Initialize student-specific session ID and conversation if not exists
    if 'student_session_id' not in st.session_state:
        # Try to auto-load most recent conversation first
        if database_available and student_id:
            from database import get_user_chat_conversations, create_chat_conversation
            db = get_db()
            if db:
                try:
                    # Get existing conversations
                    existing_conversations = get_user_chat_conversations(
                        db, student_id=student_id, interface_type='student'
                    )
                    
                    if existing_conversations and len(existing_conversations) > 0:
                        # Auto-load most recent conversation
                        most_recent = existing_conversations[0]
                        st.session_state.student_session_id = most_recent.session_id
                        st.session_state['student_current_conversation_id'] = most_recent.id
                        
                        # Load messages from database
                        loaded_messages = load_conversation_to_session(
                            db, most_recent.session_id, 'student'
                        )
                        if loaded_messages:
                            st.session_state.student_messages = loaded_messages
                            # Show restore notification
                            restore_time = most_recent.last_activity.strftime('%d/%m/%Y %H:%M') if hasattr(most_recent, 'last_activity') and most_recent.last_activity else 'earlier'
                            st.toast(f"✓ Restored your conversation from {restore_time}", icon="🔄")
                        
                        # Log session start
                        log_student_activity(
                            db, 
                            student_id, 
                            'session_start', 
                            session_id=most_recent.session_id
                        )
                    else:
                        # Create first conversation for new students
                        st.session_state.student_session_id = str(uuid.uuid4())
                        title = f"Chat {datetime.now().strftime('%d/%m %H:%M')}"
                        new_conv = create_chat_conversation(
                            db, title=title, session_id=st.session_state.student_session_id,
                            interface_type='student', user_id=None, student_id=student_id
                        )
                        st.session_state['student_current_conversation_id'] = new_conv.id
                except Exception as e:
                    print(f"Error loading/creating student conversation: {str(e)}")
                    st.session_state.student_session_id = str(uuid.uuid4())
                finally:
                    db.close()
        else:
            st.session_state.student_session_id = str(uuid.uuid4())
    
    # Render conversation sidebar (handles conversation management)
    if database_available and student_id:
        render_conversation_sidebar('student', student_id=student_id)
    
    # Ensure student messages are initialized (will be populated by auto-load above if available)
    if 'student_messages' not in st.session_state:
        st.session_state.student_messages = []
    
    # Custom CSS for enhanced chat styling
    st.markdown("""
    <style>
    /* Subject-based color themes for chat messages */
    .subject-geography { background-color: #A7C796 !important; color: white !important; }
    .subject-history { background-color: #C9A7E3 !important; color: white !important; }
    .subject-science { background-color: #74B3A1 !important; color: white !important; }
    .subject-english { background-color: #E4C29B !important; color: #5B3E2E !important; }
    .subject-civics { background-color: #B8A7C7 !important; color: white !important; }
    
    /* Smooth scroll for chat */
    [data-testid="stChatMessageContainer"] {
        scroll-behavior: smooth;
    }
    
    /* Enhanced chat bubbles */
    [data-testid="stChatMessage"] {
        animation: fadeIn 0.4s ease-in-out;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(10px); }
        to { opacity: 1; transform: translateY(0); }
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown(f"### 🌟 Welcome, {student_name}!")
    st.markdown("*Your personal learning companion for curious minds*")
    
    # Privacy Notice Banner
    st.warning("⚠️ **Privacy Notice:** Do NOT enter personal information (name, birthdate, home/school address, or details of real people). Keep all inputs anonymous.", icon="⚠️")
    
    st.markdown("### 🔍 Montessori Learning Assistant")
    st.markdown("*Explore the universe, ask questions, and discover connections*")
    
    # Initialize student-specific subject selector
    if 'student_subjects' not in st.session_state:
        st.session_state.student_subjects = []
    
    # Get age-appropriate subjects
    from utils import map_age_to_year_levels
    year_levels = map_age_to_year_levels(age_group)
    
    # Determine subject options based on age group
    if age_group in ["3-6", "6-9", "9-12"]:
        subject_options = ["English", "Mathematics", "Science", "History", "Geography", 
                          "Art", "Music", "Technology"]
    else:  # adolescent
        subject_options = ["English", "Mathematics", "Science", "History", "Geography", 
                          "Civics", "Economics", "Technology", "Art", "Music"]
    
    selected_subjects = st.multiselect(
        "Choose your subjects:",
        subject_options,
        default=st.session_state.student_subjects,
        help="Select subjects you're interested in learning about"
    )
    st.session_state.student_subjects = selected_subjects
    
    # Year level selector for students
    if 'student_year_selector' not in st.session_state:
        st.session_state.student_year_selector = year_levels[0] if year_levels else "Year 6"
    
    if year_levels and len(year_levels) > 1:
        selected_year_level = st.selectbox(
            "What year level are you studying?",
            year_levels,
            index=year_levels.index(st.session_state.student_year_selector) if st.session_state.student_year_selector in year_levels else 0
        )
        st.session_state.student_year_selector = selected_year_level
    
    # Manage conversation history (keep last 10 messages)
    st.session_state.student_messages = manage_conversation_history(
        st.session_state.student_messages, max_history=10
    )
    
    # Display chat history
    ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🌟"
    for message in st.session_state.student_messages:
        avatar = ai_avatar if message["role"] == "assistant" else None
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
    
    # Scroll to bottom of chat after displaying messages
    if st.session_state.student_messages:
        scroll_chat_to_bottom()
    
    st.markdown("---")
    
    # File upload for students with rubric support
    st.markdown("### 📁 Upload your work for feedback (optional)")
    
    col1, col2 = st.columns(2)
    with col1:
        uploaded_work = st.file_uploader(
            "Share your writing, drawing, or project",
            type=['txt', 'pdf', 'jpg', 'png', 'docx'],
            help="Upload your work for feedback",
            key="work_upload"
        )
    
    with col2:
        uploaded_rubric = st.file_uploader(
            "Upload assessment rubric (optional)",
            type=['txt', 'pdf', 'docx'],
            help="Upload a rubric to guide the feedback",
            key="rubric_upload"
        )
    
    # Feedback button when work is uploaded
    if uploaded_work:
        # Validate file uploads for security
        work_valid, work_error = validate_file_upload(uploaded_work)
        rubric_valid, rubric_error = validate_file_upload(uploaded_rubric) if uploaded_rubric else (True, None)
        
        if not work_valid:
            st.error(f"Work file validation failed: {work_error}")
        elif not rubric_valid:
            st.error(f"Rubric file validation failed: {rubric_error}")
        else:
            if st.button("🌟 How about some feedback?", use_container_width=True, type="primary"):
                # Process work file
                work_content = ""
                with st.spinner("Reading your work..."):
                    if uploaded_work.type == "application/pdf":
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_work.read()))
                        # Handle None values from scanned/textless pages
                        extracted_pages = [page.extract_text() or "" for page in pdf_reader.pages]
                        work_content = "\n".join(extracted_pages)
                        if not work_content.strip():
                            work_content = f"[PDF uploaded: {uploaded_work.name} - appears to be scanned/image-based. Text extraction not possible.]"
                    elif uploaded_work.type in ["image/jpeg", "image/png"]:
                        try:
                            from PIL import Image
                            image = Image.open(uploaded_work)
                            work_content = f"[Student uploaded an image: {uploaded_work.name}]"
                        except:
                            work_content = "[Image could not be processed]"
                    elif uploaded_work.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = Document(io.BytesIO(uploaded_work.read()))
                        work_content = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        work_content = uploaded_work.read().decode("utf-8")
                
                # Process rubric file if uploaded
                rubric_content = ""
                if uploaded_rubric:
                    with st.spinner("Reading rubric..."):
                        if uploaded_rubric.type == "application/pdf":
                            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_rubric.read()))
                            # Handle None values from scanned/textless pages
                            extracted_pages = [page.extract_text() or "" for page in pdf_reader.pages]
                            rubric_content = "\n".join(extracted_pages)
                            if not rubric_content.strip():
                                rubric_content = f"[Rubric PDF uploaded: {uploaded_rubric.name} - appears to be scanned/image-based. Text extraction not possible.]"
                        elif uploaded_rubric.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            doc = Document(io.BytesIO(uploaded_rubric.read()))
                            rubric_content = "\n".join([para.text for para in doc.paragraphs])
                        else:
                            rubric_content = uploaded_rubric.read().decode("utf-8")
                
                # Build comprehensive feedback prompt
                feedback_prompt = f"""Please provide detailed, constructive feedback on the following student work.

STUDENT WORK:
{work_content[:2000]}

"""
                if rubric_content:
                    feedback_prompt += f"""ASSESSMENT RUBRIC:
{rubric_content[:1500]}

Please assess the work against the rubric criteria, providing specific feedback for each criterion.
"""
                
                feedback_prompt += f"""
FEEDBACK REQUIREMENTS:
1. Assess the work based on the rubric (if provided) and learning standards
2. Provide specific, actionable suggestions for improvement
3. Highlight strengths and areas of growth
4. Connect feedback to broader learning and real-world applications (Montessori approach)
5. Use encouraging, asset-based language that honors student development
6. Reference Australian Curriculum V9 standards for Year {st.session_state.get('student_year_selector', '6')}
7. Suggest next steps to deepen understanding

Keep feedback age-appropriate for {age_group} year olds."""

                # Add to conversation
                st.session_state.student_messages.append({
                    "role": "user",
                    "content": f"Please review my work: {uploaded_work.name}"
                })
                
                # Get AI feedback
                ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🤖"
                with st.chat_message("assistant", avatar=ai_avatar):
                    with st.spinner("Analyzing your work..."):
                        response = call_openai_api(
                            st.session_state.student_messages + [{"role": "system", "content": feedback_prompt}],
                            is_student=True,
                            year_level=st.session_state.get('student_year_selector', 'Year 6'),
                            subjects=selected_subjects if selected_subjects else None,
                            curriculum_type="Blended"
                        )
                        st.markdown(response)
                        st.session_state.student_messages.append({
                            "role": "assistant",
                            "content": response
                        })
                        
                        # Save to database
                        feedback_save_success = False
                        if database_available and student_id:
                            db = get_db()
                            if db:
                                try:
                                    save_conversation_message(
                                        db,
                                        session_id=st.session_state.student_session_id,
                                        interface_type='student',
                                        role='assistant',
                                        content=response,
                                        user_id=None,
                                        student_id=student_id
                                    )
                                    feedback_save_success = True
                                except Exception as e:
                                    print(f"Error saving feedback: {str(e)}")
                                    st.warning("⚠️ Unable to save feedback. Please check your connection.")
                                finally:
                                    db.close()
                        
                        # Show save confirmation if successful
                        if feedback_save_success:
                            st.toast("✓ Feedback saved", icon="💾")
    
    st.markdown("---")
    
    # Chat input - always visible regardless of file upload
    uploaded_file = uploaded_work if 'uploaded_work' in dir() else None  # Keep backward compatibility
    if prompt := st.chat_input("Ask me anything about your learning..."):
        st.session_state.student_messages.append({
            "role": "user",
            "content": prompt
        })
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Process uploaded file if present
        file_context = ""
        if uploaded_file:
            with st.spinner("Reading your file..."):
                if uploaded_file.type == "application/pdf":
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                    file_context = "\n".join([page.extract_text() for page in pdf_reader.pages])
                elif uploaded_file.type in ["image/jpeg", "image/png"]:
                    try:
                        image = Image.open(uploaded_file)
                        file_context = f"[Student uploaded an image: {uploaded_file.name}]"
                    except:
                        file_context = "[Image could not be processed]"
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = Document(io.BytesIO(uploaded_file.read()))
                    file_context = "\n".join([para.text for para in doc.paragraphs])
                else:
                    file_context = uploaded_file.read().decode("utf-8")
                
                if file_context:
                    prompt = f"{prompt}\n\n[Student's uploaded content: {file_context[:1000]}]"
        
        # Save conversation and detect curriculum keywords
        save_success = False
        detected_keywords = []  # Initialize to avoid unbound variable
        if database_available and student_id:
            db = get_db()
            if db:
                try:
                    # Save user message
                    save_conversation_message(
                        db,
                        session_id=st.session_state.student_session_id,
                        interface_type='student',
                        role='user',
                        content=prompt,
                        user_id=None,
                        student_id=student_id
                    )
                    
                    # Detect and track curriculum keywords
                    from utils import detect_trending_keywords, update_trending_keywords
                    detected_keywords = detect_trending_keywords(prompt)
                    
                    # Create extra_data with detected keywords for anonymized tracking
                    extra_data = {
                        'detected_keywords': [
                            {'subject': kw['subject'], 'keyword': kw['keyword']} 
                            for kw in detected_keywords
                        ],
                        'query_id': st.session_state.student_session_id[:8]  # Anonymous query ID
                    }
                    
                    # Log student activity with keyword data
                    log_student_activity(
                        db, 
                        student_id, 
                        'learning_question', 
                        prompt_text=prompt,
                        session_id=st.session_state.student_session_id,
                        extra_data=json.dumps(extra_data)
                    )
                    
                    # Update trending keywords
                    update_trending_keywords(
                        db, 
                        detected_keywords, 
                        st.session_state.student_session_id
                    )
                    save_success = True
                except Exception as e:
                    print(f"Error saving conversation/logging prompt: {str(e)}")
                    st.warning("⚠️ Unable to save message. Please check your connection.")
                finally:
                    db.close()
        
        # Show save confirmation if successful
        if save_success:
            st.toast("✓ Message saved", icon="💾")
        
        # Get selected subjects and year level
        selected_subjects = st.session_state.get('student_subjects', [])
        selected_year_level = st.session_state.get('student_year_selector', 'Year 6')
        
        # Get AI response with enhanced features and curriculum alignment
        ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "🤖"
        with st.chat_message("assistant", avatar=ai_avatar):
            with st.spinner("Thinking..."):
                response = call_openai_api(
                    st.session_state.student_messages,
                    is_student=True,
                    year_level=selected_year_level,
                    subjects=selected_subjects if selected_subjects else None,
                    curriculum_type="Blended",
                    use_conversation_history=True
                )
                
                # Highlight detected keywords in response
                highlighted_response = response
                if detected_keywords:
                    for kw in detected_keywords:
                        keyword = kw['keyword']
                        # Use markdown bold to highlight keywords
                        import re
                        pattern = r'\b' + re.escape(keyword) + r'\b'
                        highlighted_response = re.sub(
                            pattern, 
                            f"**{keyword}**", 
                            highlighted_response, 
                            flags=re.IGNORECASE
                        )
                
                st.markdown(highlighted_response)
                st.session_state.student_messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                # Scroll to beginning of new response
                scroll_chat_to_bottom()
                
                # Save assistant response
                assistant_save_success = False
                if database_available and student_id:
                    db = get_db()
                    if db:
                        try:
                            save_conversation_message(
                                db,
                                session_id=st.session_state.student_session_id,
                                interface_type='student',
                                role='assistant',
                                content=response,
                                user_id=None,
                                student_id=student_id
                            )
                            
                            # Log response activity
                            log_student_activity(
                                db, 
                                student_id, 
                                'ai_response', 
                                prompt_text=prompt,
                                response_text=response,
                                session_id=st.session_state.student_session_id
                            )
                            assistant_save_success = True
                        except Exception as e:
                            print(f"Error saving conversation/logging response: {str(e)}")
                            st.warning("⚠️ Unable to save response. Please check your connection.")
                        finally:
                            db.close()
                
                # Show save confirmation if successful
                if assistant_save_success:
                    st.toast("✓ Response saved", icon="💾")
    
    # Add scroll to top button
    add_scroll_to_top_button()
    
def show_student_dashboard_interface():
    """Student observation dashboard for educators"""
    scroll_to_top()
    
    st.markdown("### 📊 Student Dashboard")
    st.markdown("*View student learning activities, engagement patterns, and manage access*")
    
    # Create Student button at top
    if st.button("➕ Create New Student", type="primary", use_container_width=True):
        st.session_state.auth_mode = "create_student"
        st.rerun()
    
    st.markdown("---")
    
    if not database_available:
        st.warning("Student dashboard requires database connection. Please contact support.")
        return
    
    # Get educator info
    educator_id = st.session_state.get('user_id')
    
    db = get_db()
    if not db:
        st.error("Database connection not available")
        return
    
    try:
        # Get all students accessible to this educator
        from database import get_educator_accessible_students
        students = get_educator_accessible_students(db, educator_id)
        
        if not students:
            st.info("No students found. Create student accounts to get started.")
            return
        
        search_query = st.text_input("🔍 Search students by name", placeholder="Type a name to filter...", key="student_search")
        
        if search_query and search_query.strip():
            filtered_students = [s for s in students if search_query.strip().lower() in s.full_name.lower()]
        else:
            filtered_students = students
        
        if not filtered_students:
            st.info(f"No students matching '{search_query}'. Try a different name.")
            return
        
        selected_student = st.selectbox(
            f"Select Student ({len(filtered_students)} of {len(students)}):",
            filtered_students,
            format_func=lambda s: f"{s.full_name} ({s.age_group})"
        )
        
        if selected_student:
            # Check if current educator is primary educator
            is_primary_educator = selected_student.educator_id == educator_id
            
            # Display student info
            st.markdown(f"## {selected_student.full_name}")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Age Group", selected_student.age_group or "Not specified")
            with col2:
                st.metric("Primary Educator", "You" if is_primary_educator else selected_student.educator.full_name)
            with col3:
                st.metric("Status", "Active" if selected_student.is_active else "Inactive")
            
            # Tabs for different views
            tab1, tab2, tab3, tab4 = st.tabs(["📚 Learning Activity", "💬 Chat History", "🔐 Access Management", "🚨 Safety Alerts"])
            
            with tab1:
                # Get student activities
                from database import get_student_activities
                activities = get_student_activities(db, selected_student.id, limit=50)
                
                if activities:
                    # Show summary metrics
                    st.markdown("### 📈 Engagement Summary")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Total Activities", len(activities))
                    with col2:
                        question_count = len([a for a in activities if a.activity_type == 'learning_question'])
                        st.metric("Questions Asked", question_count)
                    with col3:
                        # Get latest activity timestamp
                        if activities:
                            latest = activities[0].created_at.strftime("%d/%m/%Y")
                            st.metric("Last Activity", latest)
                    
                    # Activity timeline
                    st.markdown("### 📅 Recent Activity")
                    for activity in activities[:10]:  # Show last 10
                        # Build descriptive title with preview and subject
                        activity_preview = ""
                        subject_tag = ""
                        
                        # Get subject from extra_data if available
                        if activity.extra_data:
                            try:
                                extra = json.loads(activity.extra_data)
                                if 'detected_keywords' in extra and extra['detected_keywords']:
                                    # Get first detected subject
                                    subject_tag = f"[{extra['detected_keywords'][0]['subject']}]"
                            except:
                                pass
                        
                        # Create preview from prompt text
                        if activity.prompt_text:
                            # Clean and truncate preview
                            preview_text = activity.prompt_text.replace('\n', ' ').strip()
                            if len(preview_text) > 70:
                                activity_preview = preview_text[:70] + "..."
                            else:
                                activity_preview = preview_text
                        else:
                            activity_preview = "AI Response"
                        
                        # Format timestamp with date and time
                        timestamp = activity.created_at.strftime('%d/%m/%Y %I:%M %p')
                        
                        # Build expander title
                        icon = "💭" if activity.activity_type == 'learning_question' else "🤖"
                        expander_title = f"{icon} {subject_tag} {activity_preview} • {timestamp}"
                        
                        with st.expander(expander_title):
                            if activity.prompt_text:
                                st.markdown(f"**Student:**")
                                st.markdown(activity.prompt_text)
                            if activity.response_text:
                                st.markdown(f"**AI:**")
                                st.markdown(activity.response_text)
                            if activity.extra_data:
                                try:
                                    extra = json.loads(activity.extra_data)
                                    if 'detected_keywords' in extra:
                                        keywords = [f"{kw['subject']}: {kw['keyword']}" for kw in extra['detected_keywords']]
                                        st.markdown(f"**Detected Topics:** {', '.join(keywords)}")
                                except:
                                    pass
                else:
                    st.info("No activities recorded yet for this student.")
            
            with tab2:
                # Chat History Tab - View student chats filtered by subject
                st.markdown("### 💬 Student Chat History")
                
                from database import get_available_subjects, get_filtered_student_chats
                
                # Subject filter
                subjects = get_available_subjects()
                col1, col2 = st.columns([3, 1])
                with col1:
                    selected_subject = st.selectbox(
                        "Filter by Subject:",
                        ["All Subjects"] + subjects,
                        key=f"subject_filter_{selected_student.id}"
                    )
                with col2:
                    st.markdown("<div style='padding-top: 28px;'></div>", unsafe_allow_html=True)
                    st.caption(f"Total conversations shown below")
                
                # Get filtered chats
                subject_filter = None if selected_subject == "All Subjects" else selected_subject
                chats = get_filtered_student_chats(
                    db, 
                    educator_id=educator_id,
                    student_id=selected_student.id, 
                    subject_tag=subject_filter
                )
                
                if chats:
                    st.caption(f"{len(chats)} conversation(s)")
                    
                    # Display all chats chronologically (no subject grouping)
                    for chat in chats:
                        # Simple expander title: subject tag + title + date
                        subject_label = f"[{chat.subject_tag}] " if chat.subject_tag else ""
                        expander_title = f"{subject_label}{chat.title} • {chat.updated_at.strftime('%d/%m/%Y')}"
                        
                        with st.expander(expander_title):
                            # Load and display conversation messages naturally
                            from database import load_conversation_to_session
                            messages = load_conversation_to_session(db, chat.session_id, 'student')
                            
                            if messages:
                                for msg in messages:
                                    if msg['role'] == 'user':
                                        st.markdown("**Student:**")
                                    else:
                                        st.markdown("**AI:**")
                                    st.markdown(msg["content"])
                                    st.divider()
                            else:
                                st.info("No messages in this conversation yet.")
                else:
                    st.info(f"No chat conversations found{' for ' + selected_subject if selected_subject != 'All Subjects' else ''}.")
            
            with tab3:
                st.markdown("### 🔐 Access Management")
                
                if is_primary_educator:
                    st.info("As the primary educator, you have full access to this student's data.")
                    
                    # Grant access to other educators
                    st.markdown("### Share Access")
                    from database import get_all_educators, grant_educator_access, get_student_access_educators, is_institution_enforcement_on, User
                    
                    all_educators = get_all_educators(db)
                    
                    # Filter educators by institution if enforcement is ON
                    enforcement_on = is_institution_enforcement_on(db)
                    if enforcement_on:
                        current_educator = db.query(User).filter(User.id == educator_id).first()
                        if current_educator and current_educator.institution_name:
                            # Only show educators from same institution
                            other_educators = [e for e in all_educators 
                                             if e.id != educator_id 
                                             and e.institution_name == current_educator.institution_name]
                        else:
                            other_educators = [e for e in all_educators if e.id != educator_id]
                    else:
                        other_educators = [e for e in all_educators if e.id != educator_id]
                    
                    if other_educators:
                        selected_educator = st.selectbox(
                            "Grant access to educator:",
                            other_educators,
                            format_func=lambda e: f"{e.full_name} ({e.email})"
                        )
                        
                        if st.button("Grant Access"):
                            success, error_msg = grant_educator_access(db, selected_educator.id, selected_student.id, educator_id)
                            if success:
                                st.success(f"✅ Access granted to {selected_educator.full_name}")
                                st.rerun()
                            else:
                                st.error(f"❌ {error_msg or 'Failed to grant access'}")
                    
                    # Show current access list
                    access_list = get_student_access_educators(db, selected_student.id)
                    if access_list:
                        st.markdown("### Current Access")
                        for educator in access_list:
                            st.markdown(f"- {educator.full_name} ({educator.email})")
                    
                    # Danger zone - remove student
                    st.markdown("---")
                    st.markdown("### ⚠️ Danger Zone")
                    with st.expander("Remove Student Account"):
                        st.warning("This will permanently delete the student account and all associated data (activities, conversations, files, consent records). This action cannot be undone.")
                        deletion_reason = st.text_input("Reason for deletion (optional):", key="deletion_reason",
                                                       placeholder="e.g., Student left school, Parent request")
                        confirm_text = st.text_input("Type the student's username to confirm:", key="confirm_delete")
                        if st.button("Permanently Delete Student Account", type="primary"):
                            if confirm_text == selected_student.username:
                                from database import delete_student_and_data
                                result = delete_student_and_data(db, selected_student.id, educator_id, deletion_reason or None)
                                if result.get('success'):
                                    st.success(f"✅ Successfully removed {selected_student.full_name}'s account and all data")
                                    st.info(f"Deleted: {result.get('activities_deleted', 0)} activities, "
                                           f"{result.get('conversations_deleted', 0)} conversations, "
                                           f"{result.get('consents_deleted', 0)} consent records")
                                    st.caption("A permanent audit record of this deletion has been created.")
                                    st.rerun()
                                else:
                                    st.error(f"❌ {result.get('error', 'Failed to remove student account. Please try again.')}")
                else:
                    st.info("You have shared access to this student. Contact the primary educator for access management.")
            
            with tab4:
                st.markdown("### 🚨 Safety Alerts")
                st.markdown("*Review alerts for concerning content or student-reported concerns*")
                
                if is_primary_educator:
                    from database import SafetyAlert, review_safety_alert
                    import json as json_module
                    
                    # Get safety alerts for this student
                    alerts = db.query(SafetyAlert).filter(
                        SafetyAlert.student_id == selected_student.id
                    ).order_by(SafetyAlert.created_at.desc()).all()
                    
                    if alerts:
                        # Count pending alerts
                        pending_count = len([a for a in alerts if a.status == 'pending'])
                        reviewed_count = len([a for a in alerts if a.status != 'pending'])
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("Pending Review", pending_count, 
                                     delta="Needs attention" if pending_count > 0 else None,
                                     delta_color="inverse" if pending_count > 0 else "off")
                        with col2:
                            st.metric("Reviewed", reviewed_count)
                        
                        st.markdown("---")
                        
                        # Show alerts
                        for alert in alerts:
                            severity_colors = {
                                'high': '🔴',
                                'medium': '🟡',
                                'low': '🟢'
                            }
                            severity_icon = severity_colors.get(alert.severity, '⚪')
                            
                            status_badge = "🔔 **PENDING**" if alert.status == 'pending' else f"✅ {alert.status.title()}"
                            
                            with st.expander(f"{severity_icon} {alert.alert_type.replace('_', ' ').title()} - {alert.created_at.strftime('%d/%m/%Y %H:%M')} {status_badge}"):
                                st.markdown(f"**Severity:** {alert.severity.title()}")
                                st.markdown(f"**Type:** {alert.alert_type.replace('_', ' ').title()}")
                                
                                if alert.trigger_text:
                                    st.markdown("**Content:**")
                                    st.text_area("", value=alert.trigger_text, disabled=True, key=f"alert_text_{alert.id}")
                                
                                if alert.matched_keywords:
                                    try:
                                        keywords = json_module.loads(alert.matched_keywords)
                                        st.markdown(f"**Detected Keywords:** {', '.join(keywords)}")
                                    except:
                                        pass
                                
                                if alert.context:
                                    st.markdown(f"**Context:** {alert.context}")
                                
                                if alert.status == 'pending':
                                    st.markdown("---")
                                    st.markdown("**Take Action:**")
                                    col1, col2, col3 = st.columns(3)
                                    with col1:
                                        if st.button("✅ Mark Reviewed", key=f"review_{alert.id}"):
                                            review_safety_alert(db, alert.id, educator_id, 'reviewed')
                                            st.success("Alert marked as reviewed")
                                            st.rerun()
                                    with col2:
                                        if st.button("📋 Mark Actioned", key=f"action_{alert.id}"):
                                            review_safety_alert(db, alert.id, educator_id, 'actioned')
                                            st.success("Alert marked as actioned")
                                            st.rerun()
                                    with col3:
                                        if st.button("❌ Dismiss", key=f"dismiss_{alert.id}"):
                                            review_safety_alert(db, alert.id, educator_id, 'dismissed')
                                            st.info("Alert dismissed")
                                            st.rerun()
                                    
                                    notes = st.text_area("Review notes:", key=f"notes_{alert.id}", 
                                                        placeholder="Add notes about your follow-up actions...")
                                    if st.button("Save Notes", key=f"save_notes_{alert.id}"):
                                        review_safety_alert(db, alert.id, educator_id, alert.status, notes)
                                        st.success("Notes saved")
                                else:
                                    if alert.review_notes:
                                        st.markdown(f"**Review Notes:** {alert.review_notes}")
                                    if alert.reviewed_at:
                                        st.caption(f"Reviewed on {alert.reviewed_at.strftime('%d/%m/%Y %H:%M')}")
                    else:
                        st.info("No safety alerts for this student. This is good news!")
                        st.markdown("""
                        **Safety alerts are generated when:**
                        - Concerning content is detected in student messages (self-harm, bullying indicators)
                        - A student uses the "Need to talk to someone?" feature to report a concern
                        """)
                else:
                    st.info("Safety alerts are only visible to the primary educator.")
    
    except Exception as e:
        st.error(f"Error loading student data: {str(e)}")
        print(f"Dashboard error: {str(e)}")
        # CRITICAL: Rollback transaction on error to prevent "transaction aborted" errors
        if db:
            try:
                db.rollback()
            except Exception as rollback_error:
                print(f"Rollback error: {str(rollback_error)}")
    
    finally:
        if db:
            db.close()
    
    # Add scroll to top button
    add_scroll_to_top_button()


def show_great_story_interface():
    """Montessori Great Story creator interface for educators"""
    scroll_to_top()
    
    st.markdown("### 📖 Montessori Great Story Creator")
    st.markdown("*Develop inspiring cosmic education stories that spark imagination and curiosity*")
    
    if not database_available:
        st.info("Stories will not be saved without database connection. Your stories will be available during this session only.")
    
    educator_id = st.session_state.get('user_id')
    
    # Tabs for creating new stories and viewing saved stories
    tab1, tab2 = st.tabs(["✨ Create New Story", "📚 My Saved Stories"])
    
    with tab1:
        st.markdown("### Theme or Topic")
        st.caption("Enter a theme or topic to develop an immersive historical narrative Great Story")
        
        # Theme/topic input
        theme = st.text_input("Story Theme or Topic:", placeholder="e.g., Life in Medieval Europe, The Magna Carta, Ancient Trade Routes")
        
        # Year level selector for historical narratives (Years 7-9)
        year_level = st.selectbox(
            "Target Year Level:",
            ["Year 7", "Year 8", "Year 9"],
            index=1,
            help="Select the year level to adjust narrative complexity, vocabulary, and reflection prompts"
        )
        
        # Map year level to age group for database storage
        age_group_map = {"Year 7": "12-15", "Year 8": "12-15", "Year 9": "12-15"}
        age_group = age_group_map.get(year_level, "12-15")
        
        # Story style info
        st.info("📚 **Historical Narrative Style**: Stories follow our immersive, observational approach with quiet warmth, concrete sensory detail, ADHD-friendly paragraphs, and subtle reflection prompts embedded throughout. Every story includes a Historical Note for factual verification.")
        
        # Story development prompts
        st.markdown("### Story Development Assistance")
        
        cols = st.columns(2)
        with cols[0]:
            if st.button("Generate Story Outline", use_container_width=True):
                if theme:
                    from utils import get_great_story_system_prompt
                    system_prompt = f"""IMPORTANT: Always use British English spelling and conventions (colour, organisation, analyse, centre, programme, etc.) in all responses.

You are an AI historian and storyteller specialising in immersive historical narratives for {year_level} students (ages 12-15).

Create a detailed story outline for the theme: "{theme}"

The outline should:
- Ground the reader in time and place through sensory observation
- Identify key historical figures AND ordinary people to feature
- Plan 3-5 subtle reflection prompts to embed throughout
- Note specific historical details to verify for accuracy
- Structure: Opening (sensory grounding) → Development (human-scale narrative) → Agency moments → Closing reflection
- Follow our historical narrative style: quiet, observational, warm but unsentimental

Provide a structured outline with:
1. Opening scene (sensory details)
2. Key story beats with historical context
3. Suggested reflection prompts
4. Closing that links past to present
5. Historical details to verify"""
                    
                    with st.spinner("Creating story outline..."):
                        outline = call_openai_api(
                            [{"role": "user", "content": system_prompt}],
                            curriculum_type="Montessori",
                            interface_type="great_stories"
                        )
                        st.session_state.story_outline = outline
                        st.markdown(outline)
                else:
                    st.warning("Please enter a theme or topic first")
        
        with cols[1]:
            if st.button("Generate Full Story", use_container_width=True):
                if theme:
                    from utils import get_great_story_system_prompt
                    outline_context = st.session_state.get('story_outline', '')
                    
                    # Use the comprehensive training pack system prompt
                    system_prompt = get_great_story_system_prompt(
                        year_level=year_level,
                        theme=theme,
                        outline_context=outline_context
                    )
                    
                    with st.spinner("Crafting your Great Story..."):
                        story = call_openai_api(
                            [{"role": "user", "content": system_prompt}],
                            curriculum_type="Montessori",
                            max_tokens=3500,
                            interface_type="great_stories"
                        )
                        st.session_state.generated_story = story
                        st.markdown("### 📖 Your Great Story")
                        st.markdown(story)
                        
                        # Factual accuracy reminder
                        st.caption("💡 **Tip**: Check the Historical Note section at the end for key facts to verify. The AI has been instructed to prioritise accuracy, but we recommend cross-checking specific dates, names, and events.")
                        
                        # Save story option
                        if database_available and educator_id:
                            story_title = st.text_input("Story Title:", value=theme)
                            if st.button("💾 Save Story"):
                                db = get_db()
                                if db:
                                    try:
                                        from database import create_great_story
                                        create_great_story(db, educator_id, story_title, theme, story, age_group)
                                        st.success("✅ Story saved successfully!")
                                    except Exception as e:
                                        st.error(f"Error saving story: {str(e)}")
                                    finally:
                                        db.close()
                else:
                    st.warning("Please enter a theme or topic first")
    
    with tab2:
        if database_available and educator_id:
            db = get_db()
            if db:
                try:
                    from database import get_educator_great_stories
                    stories = get_educator_great_stories(db, educator_id)
                    
                    if stories:
                        st.markdown("### 📚 Your Saved Stories")
                        for story in stories:
                            with st.expander(f"{story.title} ({story.age_group}) - {story.created_at.strftime('%Y-%m-%d')}"):
                                st.markdown(f"**Theme:** {story.theme or 'Not specified'}")
                                st.markdown(story.content)
                                
                                # Delete option
                                if st.button(f"🗑️ Delete", key=f"delete_story_{story.id}"):
                                    from database import delete_great_story
                                    if delete_great_story(db, story.id):
                                        st.success("Story deleted")
                                        st.rerun()
                    else:
                        st.info("No saved stories yet. Create your first Great Story!")
                except Exception as e:
                    st.error(f"Error loading stories: {str(e)}")
                finally:
                    db.close()
        else:
            st.info("Database connection required to view saved stories.")


def show_planning_notes_interface():
    """Notes and planning workspace for educators"""
    scroll_to_top()
    
    st.markdown("### 📝 Planning Notes & Workspace")
    st.markdown("*Your personal workspace for planning, notes, resources, and materials organization*")
    
    if not database_available:
        st.info("Notes will not be saved without database connection. Your notes will be available during this session only.")
    
    educator_id = st.session_state.get('user_id')
    
    # Tabs for creating new notes and viewing saved notes
    tab1, tab2 = st.tabs(["📝 Active Workspace", "📂 My Saved Notes"])
    
    with tab1:
        # Note selection or creation
        st.markdown("### Select or Create Note")
        
        db = get_db()
        existing_notes = []
        if db and educator_id:
            try:
                from database import get_educator_planning_notes
                existing_notes = get_educator_planning_notes(db, educator_id)
            except Exception as e:
                print(f"Error loading notes: {str(e)}")
            finally:
                db.close()
        
        # Create new or select existing
        col1, col2 = st.columns([3, 1])
        with col1:
            if existing_notes:
                note_options = ["-- Create New Note --"] + [f"{note.title}" for note in existing_notes]
                
                # Determine default index based on active_note_id
                default_index = 0
                active_note_id = st.session_state.get('active_note_id')
                if active_note_id:
                    for idx, note in enumerate(existing_notes):
                        if note.id == active_note_id:
                            default_index = idx + 1  # +1 because "Create New Note" is at index 0
                            break
                
                selected_note_name = st.selectbox("Select Note:", note_options, index=default_index)
                
                if selected_note_name != "-- Create New Note --":
                    # Load selected note
                    selected_note = next((n for n in existing_notes if n.title == selected_note_name), None)
                    if selected_note:
                        st.session_state.active_note_id = selected_note.id
                        st.session_state.note_title = selected_note.title
                        st.session_state.note_content = selected_note.content
                        st.session_state.note_materials = selected_note.materials or ""
                        
                        # Optionally display images
                        if hasattr(selected_note, 'image_data') and selected_note.image_data:
                            try:
                                image = Image.open(io.BytesIO(selected_note.image_data))
                                st.image(image, caption="Attached Image", use_column_width=True)
                            except:
                                pass
                else:
                    # Reset for new note
                    st.session_state.active_note_id = None
                    st.session_state.note_title = ""
                    st.session_state.note_content = ""
                    st.session_state.note_materials = ""
            else:
                st.info("No existing notes. Create your first note below!")
        
        with col2:
            if st.button("🆕 New Note", use_container_width=True):
                st.session_state.active_note_id = None
                st.session_state.note_title = ""
                st.session_state.note_content = ""
                st.session_state.note_materials = ""
                st.rerun()
        
        st.markdown("---")
        
        # Note editor
        title = st.text_input("Note Title:", value=st.session_state.get('note_title', ''), placeholder="Enter note title...")
        
        # Note content
        content = st.text_area(
            "Notes Content:", 
            value=st.session_state.get('note_content', ''),
            height=300,
            placeholder="Write your planning notes, ideas, observations..."
        )
        
        # Materials list
        st.markdown("### 🧰 Materials & Resources")
        materials = st.text_area(
            "Materials List:",
            value=st.session_state.get('note_materials', ''),
            height=150,
            placeholder="List materials, resources, or links needed..."
        )
        
        # Image attachment
        uploaded_image = st.file_uploader("Attach Image:", type=['jpg', 'jpeg', 'png'])
        
        # Save button
        if st.button("💾 Save Note", use_container_width=True, type="primary"):
            if title and content:
                # Validate image if uploaded
                if uploaded_image:
                    img_valid, img_error = validate_file_upload(uploaded_image)
                    if not img_valid:
                        st.error(f"Image validation failed: {img_error}")
                        st.stop()
                
                db = get_db()
                if db and educator_id:
                    try:
                        from database import create_planning_note, update_planning_note
                        
                        # Process image if uploaded
                        image_data = None
                        if uploaded_image:
                            image_data = uploaded_image.read()
                        
                        if st.session_state.get('active_note_id'):
                            # Update existing note
                            update_planning_note(
                                db, 
                                st.session_state.active_note_id,
                                title=title, 
                                content=content,
                                images=image_data,
                                materials=materials
                            )
                            st.success("✅ Note updated successfully!")
                        else:
                            # Save new note
                            note = create_planning_note(
                                db, 
                                educator_id, 
                                title, 
                                content,
                                images=image_data,
                                materials=materials
                            )
                            st.session_state.active_note_id = note.id
                            st.success("✅ Note saved successfully!")
                        
                        st.session_state.note_title = title
                        st.session_state.note_content = content
                        st.session_state.note_materials = materials
                    except Exception as e:
                        st.error(f"Error saving note: {str(e)}")
                    finally:
                        db.close()
            else:
                st.warning("Please enter both title and content")
    
    with tab2:
        if database_available and educator_id:
            db = get_db()
            if db:
                try:
                    from database import get_educator_planning_notes, delete_planning_note
                    notes = get_educator_planning_notes(db, educator_id)
                    
                    if notes:
                        st.markdown("### 📂 Your Saved Notes")
                        for note in notes:
                            with st.expander(f"{note.title} - {note.created_at.strftime('%Y-%m-%d')}"):
                                st.markdown(note.content)
                                
                                if note.materials:
                                    st.markdown("**Materials:**")
                                    st.markdown(note.materials)
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.button(f"✏️ Edit", key=f"edit_{note.id}"):
                                        # Load note data into session state for editing
                                        st.session_state.active_note_id = note.id
                                        st.session_state.note_title = note.title
                                        st.session_state.note_content = note.content
                                        st.session_state.note_materials = note.materials or ""
                                        st.session_state.auth_mode = "planning_notes"
                                        st.rerun()
                                with col2:
                                    if st.button(f"🗑️ Delete", key=f"delete_{note.id}"):
                                        if delete_planning_note(db, note.id):
                                            st.success("Note deleted")
                                            st.rerun()
                    else:
                        st.info("No saved notes yet.")
                except Exception as e:
                    st.error(f"Error loading notes: {str(e)}")
                finally:
                    db.close()


def show_privacy_policy():
    """Terms and Conditions page with Terms of Use and Privacy Policy"""
    st.markdown("### 📋 Terms and Conditions")
    st.markdown("*Last Updated: October 2025*")
    
    # Add tabs for Terms of Use and Privacy Policy
    tab1, tab2 = st.tabs(["Terms of Use", "Privacy Policy"])
    
    with tab1:
        st.markdown("""
    ## Terms of Use
    
    **Last Updated:** October 2025
    
    Welcome to **Guide — Your digital prepared environment** ("**Guide**," "**we**," "**us**," or "**our**"), a product of **Auxpery** ("**Auxpery**," "**Company**").
    
    By creating an account or using Guide, you agree to these Terms of Use and our Privacy Policy. Please read them carefully. If you do not agree, you must not use the platform.
    
    ---
    
    ## 1. Purpose and Scope
    
    Guide is an educational platform designed to support Montessori and Australian Curriculum–aligned teaching and learning.
    
    These Terms apply to all users, including:
    
    - **Educators** – individuals creating lessons, managing students, or accessing educator tools
    - **Students** – individuals using learning features under educator or parental supervision
    - **Schools or Organisations** – entities managing multiple educator or student accounts
    
    Use of Guide constitutes agreement to these Terms.
    
    ---
    
    ## 2. Account Registration and Responsibility
    
    ### Educators
    
    - Must provide accurate information when registering.
    - Are responsible for maintaining confidentiality of login credentials.
    - Must ensure all student accounts under their supervision are used appropriately and with parental consent (if applicable).
    
    ### Students
    
    - Must only use Guide under supervision of an educator or parent/guardian.
    - Should never share personal or identifying information in prompts, messages, or uploads.
    
    We reserve the right to suspend or terminate accounts that breach these Terms or pose security or safety risks.
    
    ---
    
    ## 3. Acceptable Use
    
    You agree to use Guide only for lawful educational purposes and not to:
    
    - Upload or share offensive, discriminatory, or unlawful content
    - Misuse AI functions to generate or share harmful material
    - Attempt to access other users' data or interfere with the system
    - Use the platform for commercial, advertising, or non-educational purposes
    - Reverse-engineer or extract any source code, AI model, or backend data
    
    We may monitor activity to ensure platform integrity and educational safety.
    
    ---
    
    ## 4. Intellectual Property
    
    All content, design elements, and software comprising Guide — including logos, interface, and educational tools — are the property of **Auxpery** or its licensors and are protected by copyright, trademark, and other laws.
    
    You may use materials within Guide **solely for educational, non-commercial use**.
    
    You may not reproduce, modify, distribute, or create derivative works without written permission.
    
    Any user-generated content (such as lesson plans or educator resources) remains your intellectual property, but by uploading it, you grant Auxpery a **non-exclusive, royalty-free licence** to host, display, and use it within the platform.
    
    ---
    
    ## 5. AI-Generated Content
    
    Guide uses OpenAI's language models to generate responses, lesson ideas, and learning materials.
    
    - AI-generated content is intended for **educational guidance only**.
    - Educators remain responsible for reviewing and adapting AI output before classroom use.
    - We do not guarantee factual accuracy or curriculum compliance of every AI response.
    - Users must not rely on AI-generated content for legal, medical, or financial advice.
    
    ---
    
    ## 6. Payments and Subscriptions
    
    Guide currently offers subscription-based access.
    
    **Educator accounts:**
    
    - Include unlimited student accounts per educator.
    - Subscription fees are billed monthly or annually, as displayed at checkout.
    - Payment is processed securely via an approved provider (e.g., Stripe or PayPal).
    
    **Cancellations:**
    
    - You may cancel at any time. Access continues until the end of the paid period.
    - Refunds are provided only where required by Australian Consumer Law.
    
    **Changes to pricing or plans:**
    
    - We may update fees or plans with 30 days' prior notice. Continued use after notice means acceptance of the new terms.
    
    ---
    
    ## 7. Privacy and Data Protection
    
    Your use of Guide is governed by our **Privacy Policy**, which explains how we collect, store, and use your data in compliance with the Australian Privacy Act 1988.
    
    By using Guide, you consent to the collection and overseas transfer of data as described in that policy.
    
    ---
    
    ## 8. Third-Party Services
    
    Guide integrates third-party services such as OpenAI (for AI assistance) and Replit (for hosting).
    
    We do not control these services and are not responsible for their content, availability, or data handling.
    
    Use of these features is at your discretion and governed by their respective terms and privacy policies.
    
    ---
    
    ## 9. Availability and Service Changes
    
    We aim for continuous service but do not guarantee uninterrupted access.
    
    We may modify, suspend, or discontinue parts of the platform (including beta features) at any time for maintenance, updates, or improvement.
    
    We will give notice of significant service changes where practical.
    
    ---
    
    ## 10. Disclaimer and Limitation of Liability
    
    Guide and all related content are provided "as is."
    
    To the extent permitted by law:
    
    - We make no warranties about completeness, reliability, or suitability of AI-generated or user-contributed content.
    - Auxpery and its partners are not liable for indirect, incidental, or consequential losses arising from use or inability to use the platform.
    
    Nothing in these Terms limits your rights under the **Australian Consumer Law (ACL)**.
    
    ---
    
    ## 11. Termination of Use
    
    We may suspend or terminate access if:
    
    - These Terms are violated
    - Required by law or regulation
    - Platform integrity or security is at risk
    
    Upon termination, your right to use Guide ends immediately.
    
    Data will be handled according to our **Privacy Policy** and data retention schedule.
    
    ---
    
    ## 12. Children and Parental Consent
    
    For students under 18:
    
    - Accounts must be created and supervised by an educator or guardian.
    - Parents and guardians may review or request deletion of their child's data at any time.
    - Guide provides safeguards and in-app warnings to discourage personal data sharing.
    
    ---
    
    ## 13. Changes to These Terms
    
    We may update these Terms periodically.
    
    Material updates will be communicated via email or in-app notice.
    
    Your continued use of Guide after changes take effect constitutes acceptance of the new Terms.
    
    ---
    
    ## 14. Contact Us
    
    If you have any questions about these Terms or your use of Guide, please contact:
    
    **Auxpery**
    
    📧 Email: guide@auxpery.com.au
    
    ---
    
    **© 2025 Auxpery — Gentle Technology for Thoughtful Education**
    """)
    
    with tab2:
        st.markdown("""
    ## Privacy Policy
    
    **Last Updated:** December 2025
    
    ### 1. Introduction
    
    Welcome to **Guide — Your digital prepared environment** ("**Guide**," "**we**," "**us**," or "**our**").
    
    Guide is provided by **Auxpery** and designed to support educators and students in Montessori and Australian Curriculum–aligned learning environments.
    
    This Privacy Policy explains how we collect, use, store, and protect your personal information in accordance with the **Australian Privacy Act 1988** and the **Australian Privacy Principles (APPs)**.
    
    By using Guide, you consent to the practices described below.
    
    ---
    
    ### 2. Information We Collect (APP 3)
    
    **From Educators:**
    
    - Full name
    - Email address
    - Securely hashed password (never stored in plain text)
    - Usage data (prompts submitted, subjects taught, year levels)
    - Conversation history with our AI assistant
    - Uploaded documents and files for curriculum analysis
    
    **From Students:**
    
    - Full name and username
    - Securely hashed password
    - Age group (3–6, 6–9, 9–12, 12–15 years)
    - Learning activities and questions
    - Conversation history with our AI tutor
    - Uploaded files for learning support
    
    **Automatically Collected:**
    
    - Session identifiers (anonymised for students)
    - Curriculum keywords in queries
    - Usage timestamps
    - Aggregate analytics and trending topics
    
    **We Do NOT Collect:**
    
    - Home addresses
    - Birthdates
    - Phone numbers
    - Details about real people
    - Any other sensitive personal information
    
    We actively instruct students **not** to share personal or identifying information in any message or upload.
    
    ---
    
    ### 3. How We Use Your Information (APP 6)
    
    **For Educators:**
    
    - Provide AI-assisted lesson planning and curriculum alignment
    - Generate educational resources
    - Track usage analytics to improve performance
    - Save conversation history for lesson continuity
    - Enable educator collaboration features
    
    **For Students:**
    
    - Deliver personalised learning support
    - Track learning progress and engagement
    - Adjust AI responses to age and curriculum level
    - Allow educators to monitor student activity
    
    **For All Users:**
    
    - Authenticate and secure accounts
    - Improve our AI models and services
    - Comply with legal obligations
    - Communicate important service updates
    
    ---
    
    ### 4. Overseas Disclosure (APP 8)
    
    Some data is processed outside Australia to deliver Guide's services.
    
    **OpenAI (United States):**
    
    We use OpenAI's models to power the AI assistant. When you submit a question or upload content, your input and limited context are transmitted securely to OpenAI's servers in the U.S. for processing.
    
    **Data sent may include:**
    
    - Prompts and questions
    - Recent conversation history (last 10 messages)
    - Curriculum context
    - Uploaded educational content
    
    **Privacy note:**
    
    - OpenAI operates under U.S. law, not the Australian Privacy Act.
    - Their data handling is governed by the OpenAI Privacy Policy.
    - Data may be used to improve AI performance (per OpenAI's terms).
    
    By using Guide, you consent to this overseas data transfer. You may choose to avoid submitting sensitive information.
    
    **Replit (United States):**
    
    Our database and infrastructure are hosted on Replit's platform in the U.S., which stores and processes account data.
    
    ---
    
    ### 5. Data Security (APP 11)
    
    **Technical safeguards:**
    
    - All passwords hashed with bcrypt encryption
    - SSL/TLS encryption for database connections
    - Role-based access control (educator/student separation)
    - Secure session management via Streamlit
    
    **Organisational safeguards:**
    
    - Limited staff access to personal data
    - Periodic security reviews
    - Incident response procedures
    
    While we implement strong protection, no system is entirely secure. Transmission of data over the internet carries some inherent risk.
    
    ---
    
    ### 6. Data Retention (APP 11.2)
    
    **Retention periods:**
    
    In accordance with Australian education record-keeping requirements:
    
    - **Student records:** Retained for 7 years from last activity (aligned with Australian school baseline requirements)
    - **Conversation history:** Retained for 7 years from creation
    - **Uploaded files:** Retained for 7 years, then automatically purged
    - **Educator accounts:** Kept while active, deleted after 7 years of inactivity
    - **Child safety records:** Retained for 25 years (extended retention for legal compliance)
    - **Audit logs:** Retained permanently for security and compliance purposes
    
    **Session Security:**
    
    - Student sessions automatically timeout after 2 hours of inactivity
    - Educator sessions timeout after 30 minutes of inactivity to protect student data access
    
    **Deletion:**
    
    - Educators may delete student accounts and all associated data at any time
    - Parents/guardians may request deletion of their child's data at any time
    - Deletion requests are processed within 7 days
    - A permanent audit record of the deletion is maintained for accountability
    - Some anonymised analytics may be retained for service improvement
    
    ---
    
    ### 7. Your Privacy Rights (APP 12 & 13)
    
    Under Australian privacy law, you may:
    
    - **Access** your personal data
    - **Request correction** of inaccurate or incomplete data
    - **Request deletion** of your account
    - **Withdraw consent** for data processing or overseas disclosure (which may limit functionality)
    
    Requests can be made via **privacy@auxpery.com.au** or through in-app support. We respond within **30 days** of a verified request.
    
    ---
    
    ### 8. Parental Rights and Student Privacy
    
    **Students under 18:**
    
    - Educators must obtain parental or guardian consent before creating accounts
    - Parents may access or request deletion of their child's data
    - Students are reminded not to share personal information
    
    **School responsibilities:**
    
    - Schools are considered **data controllers** for student data
    - Guide acts as a **data processor** on their behalf
    - Schools must ensure appropriate consent processes are in place
    
    A **Data Processing Agreement (DPA)** is available on request.
    
    ---
    
    ### 9. Child Safety Measures
    
    Guide implements comprehensive child safety features:
    
    **Content Monitoring:**
    
    - Student messages are monitored for concerning content indicators (self-harm, bullying, abuse)
    - When concerning content is detected, educators are automatically notified for follow-up
    - No automated actions are taken; educators review alerts and determine appropriate response
    
    **Student Reporting:**
    
    - Students can confidentially report concerns to their educator using the "Need to talk to someone?" feature
    - Reports are delivered privately to the supervising educator
    - Students are encouraged to speak with trusted adults for any concerns
    
    **Personal Information Protection:**
    
    - All messages sent to AI services are sanitized to remove personal information (names, emails, phone numbers, addresses)
    - Students receive prominent warnings not to share personal information
    - File uploads are validated for security and size limits
    
    **Educator Accountability:**
    
    - All educator actions on student data are logged in permanent audit trails
    - Access to student records is restricted to the supervising educator
    - Guardian consent records are stored with timestamps and attestation text
    
    ---
    
    ### 10. Cookies and Tracking
    
    Guide uses only **essential cookies** for authentication and app functionality (Streamlit session cookies).
    
    We do **not** use third-party advertising or tracking cookies.
    
    If basic analytics (e.g., traffic counts) are enabled, they are anonymised and do not identify users.
    
    ---
    
    ### 11. Updates to This Policy (APP 1)
    
    We may update this policy to reflect:
    
    - Legal or regulatory changes
    - New features or services
    - User or regulator feedback
    
    You'll be notified of material updates via email or in-app message. Continued use after changes means acceptance of the updated terms.
    
    Previous versions are available on request.
    
    ---
    
    ### 12. Contact and Complaints
    
    **Privacy Officer**
    
    📧 Email: privacy@auxpery.com.au
    
    ⏱ Response within 30 days
    
    **To lodge a complaint:**
    
    1. Email our Privacy Officer describing your concern.
    2. We'll acknowledge receipt within 5 business days.
    3. We'll investigate and reply within 30 days.
    
    If unsatisfied, you may contact:
    
    **Office of the Australian Information Commissioner (OAIC)**
    
    🌐 www.oaic.gov.au
    
    📞 1300 363 992
    
    📧 enquiries@oaic.gov.au
    
    ---
    
    ### 13. Questions
    
    If you have any questions about this Privacy Policy or how we handle your data:
    
    📧 guide@auxpery.com.au
    
    ---
    
    **© 2025 Auxpery — Gentle Technology for Thoughtful Education**
    """)
    
    # Add quick actions
    st.markdown("---")
    st.markdown("### 🔐 Privacy Actions")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Request My Data", use_container_width=True):
            st.session_state.auth_mode = "data_access"
            st.rerun()
    
    with col2:
        if st.button("✏️ Update My Information", use_container_width=True):
            st.session_state.auth_mode = "account_settings"
            st.rerun()
    
    with col3:
        if st.button("🗑️ Delete My Account", use_container_width=True):
            st.session_state.auth_mode = "account_deletion"
            st.rerun()


def show_data_access_interface():
    """User data access interface (APP 12)"""
    st.markdown("### 📥 Access Your Data")
    st.markdown("*Request a copy of all personal information we hold about you*")
    
    user_id = st.session_state.get('user_id')
    is_student = st.session_state.get('is_student', False)
    
    st.info("Under the Australian Privacy Act 1988, you have the right to access your personal information. We will provide your data in a commonly used electronic format within 30 days.")
    
    if st.button("📥 Generate My Data Export", type="primary", use_container_width=True):
        with st.spinner("Gathering your data..."):
            db = get_db()
            if db and user_id:
                try:
                    import json
                    from datetime import datetime
                    
                    export_data = {
                        "export_date": datetime.now().isoformat(),
                        "user_type": "student" if is_student else "educator",
                        "account_info": {},
                        "conversation_history": [],
                        "activities": [],
                        "saved_content": []
                    }
                    
                    if is_student:
                        # Export student data
                        from database import Student, StudentActivity, ConversationHistory
                        student = db.query(Student).filter(Student.id == user_id).first()
                        
                        if student:
                            export_data["account_info"] = {
                                "username": student.username,
                                "full_name": student.full_name,
                                "age_group": student.age_group,
                                "created_at": student.created_at.isoformat(),
                                "is_active": student.is_active
                            }
                            
                            # Get activities
                            activities = db.query(StudentActivity).filter(
                                StudentActivity.student_id == user_id
                            ).all()
                            
                            export_data["activities"] = [
                                {
                                    "type": a.activity_type,
                                    "prompt": a.prompt_text,
                                    "response": a.response_text,
                                    "date": a.created_at.isoformat()
                                } for a in activities
                            ]
                            
                            # Get conversations
                            conversations = db.query(ConversationHistory).filter(
                                ConversationHistory.student_id == user_id
                            ).all()
                            
                            export_data["conversation_history"] = [
                                {
                                    "role": c.role,
                                    "content": c.content,
                                    "date": c.created_at.isoformat()
                                } for c in conversations
                            ]
                    
                    else:
                        # Export educator data
                        from database import User, EducatorAnalytics, ConversationHistory, LessonPlan
                        user = db.query(User).filter(User.id == user_id).first()
                        
                        if user:
                            export_data["account_info"] = {
                                "email": user.email,
                                "full_name": user.full_name,
                                "user_type": user.user_type,
                                "created_at": user.created_at.isoformat(),
                                "is_active": user.is_active
                            }
                            
                            # Get analytics
                            analytics = db.query(EducatorAnalytics).filter(
                                EducatorAnalytics.user_id == user_id
                            ).all()
                            
                            export_data["analytics"] = [
                                {
                                    "interface": a.interface_type,
                                    "subject": a.subject,
                                    "year_level": a.year_level,
                                    "prompt": a.prompt_text,
                                    "tokens": a.tokens_used,
                                    "date": a.created_at.isoformat()
                                } for a in analytics
                            ]
                            
                            # Get lesson plans
                            plans = db.query(LessonPlan).filter(
                                LessonPlan.creator_id == user_id
                            ).all()
                            
                            export_data["saved_content"] = [
                                {
                                    "type": "lesson_plan",
                                    "title": p.title,
                                    "content": p.content,
                                    "created": p.created_at.isoformat()
                                } for p in plans
                            ]
                    
                    # Create downloadable JSON
                    json_data = json.dumps(export_data, indent=2)
                    
                    st.success("✅ Data export ready!")
                    st.download_button(
                        label="📥 Download My Data (JSON)",
                        data=json_data,
                        file_name=f"guide_data_export_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json"
                    )
                    
                except Exception as e:
                    st.error(f"Error generating export: {str(e)}")
                finally:
                    db.close()
            else:
                st.error("Unable to access database")


def show_account_deletion_interface():
    """Account deletion interface (APP 12/13)"""
    
    user_id = st.session_state.get('user_id')
    is_student = st.session_state.get('is_student', False)
    user_name = st.session_state.get('user_name', '')
    
    if not is_student:
        from auth import show_account_settings
        show_account_settings()
        
        st.markdown("---")
    
    st.markdown("### 🗑️ Delete My Account")
    st.markdown("*Permanently remove your account and all associated data*")
    
    st.error("⚠️ **Warning: This action cannot be undone!**")
    
    st.markdown("""
    ### What will be deleted:
    - Your account credentials
    - All conversation history
    - All saved lesson plans, stories, and notes (educators)
    - All learning activities and progress (students)
    - All uploaded files and documents
    - All analytics and usage data
    
    ### What happens next:
    - Deletion takes effect within 7 days
    - You will receive confirmation via email (educators)
    - All data is permanently removed from our systems
    - This action cannot be reversed
    """)
    
    st.markdown("---")
    
    st.markdown("### Confirm Deletion")
    
    confirm_text = st.text_input(
        f"Type your {'username' if is_student else 'full name'} to confirm:",
        help=f"Enter exactly: {user_name}"
    )
    
    reason = st.text_area(
        "Reason for deletion (optional):",
        placeholder="Help us improve by sharing why you're leaving...",
        help="This is optional but helps us understand user needs"
    )
    
    if st.button("🗑️ Permanently Delete My Account", type="primary", use_container_width=True):
        if confirm_text == user_name:
            db = get_db()
            if db and user_id:
                try:
                    if is_student:
                        from database import delete_student
                        if delete_student(db, user_id):
                            # Log reason if provided
                            if reason:
                                # Could save to a deletion_log table for analytics
                                pass
                            
                            st.success("✅ Your account has been permanently deleted.")
                            st.info("You will be logged out in 3 seconds...")
                            
                            # Clear session
                            import time
                            time.sleep(3)
                            for key in list(st.session_state.keys()):
                                del st.session_state[key]
                            st.rerun()
                        else:
                            st.error("Failed to delete account. Please contact support.")
                    else:
                        from database import delete_educator
                        # Comprehensive deletion for educators with all related data
                        success, error_msg = delete_educator(db, user_id)
                        
                        if success:
                            st.success("✅ Your account has been permanently deleted.")
                            st.info("You will be logged out in 3 seconds...")
                            
                            import time
                            time.sleep(3)
                            for key in list(st.session_state.keys()):
                                del st.session_state[key]
                            st.rerun()
                        else:
                            st.error(f"❌ {error_msg}")
                            if error_msg and "active student" in error_msg.lower():
                                st.info("💡 **Tip:** Go to 'Create Student Account' to view and delete your students first.")
                
                except Exception as e:
                    st.error(f"Error deleting account: {str(e)}")
                finally:
                    db.close()
        else:
            st.error(f"Confirmation text does not match. Please type exactly: {user_name}")


def show_pd_expert_interface():
    """Professional Development Expert Mode - Restricted to authorized educators"""
    scroll_to_top()
    apply_chatgpt_chat_style()
    
    user_email = st.session_state.get('user_email', '')
    
    # Access control - restrict to authorized emails
    authorized_pd_emails = ["guideaichat@gmail.com", "ben@hmswairoa.net"]
    if user_email not in authorized_pd_emails:
        st.error("🔒 Access Denied")
        st.info("This Professional Development Expert Mode is restricted to authorized accounts only.")
        return
    
    # Header with special badge
    st.markdown("### 🧭 Professional Development Expert Mode")
    st.markdown("*Evidence-based PD coaching with self-learning memory and Montessori expertise*")
    
    st.markdown("---")
    
    # Info box
    with st.expander("ℹ️ About PD Expert Mode", expanded=False):
        st.markdown("""
        **This specialized mode provides:**
        
        - 📚 **Evidence-based guidance** from frameworks like Harvard Instructional Moves, Edutopia, and adult learning theory
        - 🧠 **Self-learning memory** that adapts to your professional development focus areas
        - 🎯 **Montessori-aligned** coaching with deep understanding of Prepared Adult principles
        - 🔄 **Contextual responses** based on keywords like "adult learning", "instructional coaching", "workshop design"
        
        **Structured responses include:**
        1️⃣ Summary of your question  
        2️⃣ Evidence-based insight or framework  
        3️⃣ Suggested approach or structure  
        4️⃣ Montessori connection  
        5️⃣ Next steps or reflective prompt
        """)
    
    # Document upload section for PD Expert
    st.markdown("#### 📄 Upload a Document")
    st.caption("Share PD materials, research papers, or workshop resources for analysis")
    
    pd_uploaded_document = st.file_uploader(
        "Upload document for analysis",
        type=['txt', 'pdf', 'jpg', 'png', 'docx'],
        help="Upload professional development materials for feedback and analysis",
        key="pd_document_upload"
    )
    
    # Process uploaded document
    if pd_uploaded_document:
        is_valid, error_msg = validate_file_upload(pd_uploaded_document)
        if not is_valid:
            st.error(f"File validation failed: {error_msg}")
        else:
            document_content = ""
            with st.spinner("Reading your document..."):
                if pd_uploaded_document.type == "application/pdf":
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(pd_uploaded_document.read()))
                    extracted_pages = [page.extract_text() or "" for page in pdf_reader.pages]
                    document_content = "\n".join(extracted_pages)
                    if not document_content.strip():
                        document_content = f"[PDF uploaded: {pd_uploaded_document.name} - appears to be scanned/image-based. Text extraction not possible. Please describe the content or upload a text-based version.]"
                elif pd_uploaded_document.type in ["image/jpeg", "image/png"]:
                    try:
                        from PIL import Image
                        import io
                        image = Image.open(pd_uploaded_document)
                        document_content = pytesseract.image_to_string(image)
                        if not document_content.strip():
                            document_content = f"[Image uploaded: {pd_uploaded_document.name} - visual content not extractable as text]"
                    except:
                        document_content = f"[Image uploaded: {pd_uploaded_document.name} - could not process]"
                elif pd_uploaded_document.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    import io
                    doc = Document(io.BytesIO(pd_uploaded_document.read()))
                    document_content = "\n".join([para.text for para in doc.paragraphs])
                else:
                    document_content = pd_uploaded_document.read().decode('utf-8', errors='ignore')
            
            if document_content.strip():
                st.success(f"✅ Document '{sanitize_filename(pd_uploaded_document.name)}' loaded successfully!")
                
                # Store document content in session for use in chat
                st.session_state.pd_document_content = document_content
                st.session_state.pd_document_name = sanitize_filename(pd_uploaded_document.name)
                
                # Show preview
                with st.expander("📖 Document Preview", expanded=False):
                    preview_text = document_content[:2000] + ("..." if len(document_content) > 2000 else "")
                    st.text_area("Content", preview_text, height=200, disabled=True)
                
                st.info("💡 Now ask a question about this document in the chat below!")
    
    st.markdown("---")
    
    # Initialize PD messages in session state
    if 'pd_messages' not in st.session_state:
        st.session_state.pd_messages = []
    
    # Display conversation history with avatar
    for msg in st.session_state.pd_messages:
        if msg['role'] == 'user':
            with st.chat_message("user"):
                st.markdown(msg['content'])
        else:
            with st.chat_message("assistant", avatar="assets/montessori-avatar.png"):
                st.markdown(f"**🧭 PD Expert**")
                st.markdown(msg['content'])
    
    # Chat input
    if user_prompt := st.chat_input("Ask your professional development question...", key="pd_expert_input"):
        # Display user message
        with st.chat_message("user"):
            st.markdown(user_prompt)
        
        st.session_state.pd_messages.append({
            "role": "user",
            "content": user_prompt
        })
        
        # Call PD Expert API (Python implementation)
        with st.chat_message("assistant", avatar="assets/montessori-avatar.png"):
            st.markdown("**🧭 PD Expert**")
            
            with st.spinner("Consulting expertise and memory... (this may take 30-60 seconds for comprehensive responses)"):
                from utils import call_pd_expert
                from openai import OpenAI
                import os
                
                try:
                    # Initialize OpenAI client
                    openai_client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
                    
                    # Build prompt with document context if available
                    full_prompt = user_prompt
                    if st.session_state.get('pd_document_content'):
                        doc_name = st.session_state.get('pd_document_name', 'uploaded document')
                        doc_content = st.session_state.get('pd_document_content', '')
                        # Truncate very long documents to avoid token limits
                        max_doc_chars = 15000
                        if len(doc_content) > max_doc_chars:
                            doc_content = doc_content[:max_doc_chars] + "\n\n[Document truncated due to length...]"
                        full_prompt = f"[User has uploaded a document: {doc_name}]\n\n--- DOCUMENT CONTENT ---\n{doc_content}\n--- END DOCUMENT ---\n\nUser's question: {user_prompt}"
                    
                    # Call PD Expert function
                    result = call_pd_expert(user_email, full_prompt, openai_client)
                    
                    if result.get('success'):
                        expert_response = result.get('output', '')
                        st.markdown(expert_response)
                        
                        st.session_state.pd_messages.append({
                            "role": "assistant",
                            "content": expert_response
                        })
                        
                        # Scroll to beginning of new response
                        scroll_chat_to_bottom()
                    else:
                        error_msg = result.get('error', 'Unknown error')
                        if 'Access denied' in error_msg:
                            st.error("🔒 Access denied. This feature is restricted.")
                        else:
                            st.error(f"Error: {error_msg}")
                        
                except Exception as e:
                    st.error(f"Error: {str(e)}")
    
    # Clear conversation and document buttons
    if st.session_state.pd_messages or st.session_state.get('pd_document_content'):
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear Conversation", use_container_width=True):
                st.session_state.pd_messages = []
                st.rerun()
        with col2:
            if st.session_state.get('pd_document_content'):
                if st.button("📄 Clear Document", use_container_width=True):
                    st.session_state.pd_document_content = None
                    st.session_state.pd_document_name = None
                    # Reset file uploader widget by clearing its key
                    if 'pd_document_upload' in st.session_state:
                        del st.session_state['pd_document_upload']
                    st.rerun()
    
    # Add scroll to top button
    add_scroll_to_top_button()


def show_imaginarium_interface():
    """Creative space for educators - free exploration with minimal guardrails"""
    apply_chatgpt_chat_style()
    inject_chat_auto_scroll()
    
    # Import database functions
    from database import (get_user_chat_conversations, create_chat_conversation, 
                         load_conversation_to_session, save_conversation_message)
    
    st.markdown("### ✨ Imaginarium")
    st.markdown("*Explore ideas freely – a space for imaginative thinking and open conversation*")
    
    # Get user info
    user_id = st.session_state.get('user_id')
    
    # Initialize session ID for Imaginarium if not exists
    if 'imaginarium_session_id' not in st.session_state:
        if database_available and user_id:
            db = get_db()
            if db:
                try:
                    # Auto-load most recent conversation for this interface
                    existing_conversations = get_user_chat_conversations(
                        db, user_id=user_id, interface_type='imaginarium'
                    )
                    
                    if existing_conversations and len(existing_conversations) > 0:
                        # Auto-load most recent conversation
                        most_recent = existing_conversations[0]
                        st.session_state.imaginarium_session_id = most_recent.session_id
                        st.session_state['imaginarium_current_conversation_id'] = most_recent.id
                        
                        # Load messages from database
                        loaded_messages = load_conversation_to_session(
                            db, most_recent.session_id, 'imaginarium'
                        )
                        if loaded_messages:
                            st.session_state.imaginarium_messages = loaded_messages
                            # Show restore notification
                            restore_time = most_recent.last_activity.strftime('%d/%m/%Y %H:%M') if hasattr(most_recent, 'last_activity') and most_recent.last_activity else 'earlier'
                            st.toast(f"✓ Restored your conversation from {restore_time}", icon="🔄")
                    else:
                        # Create first conversation for new users
                        st.session_state.imaginarium_session_id = str(uuid.uuid4())
                        title = f"Imaginarium {datetime.now().strftime('%d/%m %H:%M')}"
                        new_conv = create_chat_conversation(
                            db, title=title, session_id=st.session_state.imaginarium_session_id,
                            interface_type='imaginarium', user_id=user_id, student_id=None
                        )
                        st.session_state['imaginarium_current_conversation_id'] = new_conv.id
                except Exception as e:
                    print(f"Error loading/creating imaginarium conversation: {str(e)}")
                    st.session_state.imaginarium_session_id = str(uuid.uuid4())
                finally:
                    db.close()
        else:
            st.session_state.imaginarium_session_id = str(uuid.uuid4())
    
    # Render conversation sidebar (handles conversation management)
    if database_available and user_id:
        render_conversation_sidebar('imaginarium', user_id=user_id)
    
    # Ensure messages are initialized
    if 'imaginarium_messages' not in st.session_state:
        st.session_state.imaginarium_messages = []
    
    # Manage conversation history (keep last 20 exchanges for longer conversations)
    st.session_state.imaginarium_messages = manage_conversation_history(
        st.session_state.imaginarium_messages, max_history=40
    )
    
    # Check if last message needs a response
    need_response = (
        len(st.session_state.imaginarium_messages) > 0 and 
        st.session_state.imaginarium_messages[-1]["role"] == "user" and
        (len(st.session_state.imaginarium_messages) == 1 or 
         st.session_state.imaginarium_messages[-2]["role"] == "assistant")
    )
    
    # Display chat history
    ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "✨"
    for message in st.session_state.imaginarium_messages:
        avatar = ai_avatar if message["role"] == "assistant" else None
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
    
    # Scroll to bottom of chat after displaying messages
    if st.session_state.imaginarium_messages:
        scroll_chat_to_bottom()
    
    # If last message was from user, generate response
    if need_response:
        ai_avatar = "assets/montessori-avatar.png" if os.path.exists("assets/montessori-avatar.png") else "✨"
        with st.chat_message("assistant", avatar=ai_avatar):
            with st.spinner("Exploring creative possibilities..."):
                response = call_openai_api(
                    st.session_state.imaginarium_messages,
                    interface_type="imaginarium"
                )
                st.markdown(response)
                st.session_state.imaginarium_messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                # Scroll to beginning of new response
                scroll_chat_to_bottom()
                
                # Save assistant response to database with validation
                assistant_save_success = False
                if database_available and user_id:
                    session_id = st.session_state.get('imaginarium_session_id')
                    if not session_id:
                        print("ERROR: imaginarium_session_id is not set! Cannot save assistant response.")
                    else:
                        db = get_db()
                        if db:
                            try:
                                save_conversation_message(
                                    db,
                                    session_id=session_id,
                                    interface_type='imaginarium',
                                    role='assistant',
                                    content=response,
                                    user_id=user_id,
                                    student_id=None
                                )
                                assistant_save_success = True
                            except Exception as e:
                                print(f"Error saving conversation: {str(e)}")
                                st.warning("⚠️ Unable to save response. Please check your connection.")
                            finally:
                                db.close()
                
                # Show save confirmation if successful
                if assistant_save_success:
                    st.toast("✓ Response saved", icon="💾")
    
    # Chat input
    if prompt := st.chat_input("Share your ideas, questions, or creative thoughts..."):
        # Ensure session_id exists before saving (critical fix for save failures)
        if 'imaginarium_session_id' not in st.session_state or not st.session_state.imaginarium_session_id:
            st.session_state.imaginarium_session_id = str(uuid.uuid4())
            # Also create conversation record if needed
            if database_available and user_id:
                db = get_db()
                if db:
                    try:
                        title = f"Imaginarium {datetime.now().strftime('%d/%m %H:%M')}"
                        new_conv = create_chat_conversation(
                            db, title=title, session_id=st.session_state.imaginarium_session_id,
                            interface_type='imaginarium', user_id=user_id, student_id=None
                        )
                        st.session_state['imaginarium_current_conversation_id'] = new_conv.id
                    except Exception as e:
                        print(f"Error creating conversation: {str(e)}")
                    finally:
                        db.close()
        
        # Add user message to chat
        st.session_state.imaginarium_messages.append({
            "role": "user",
            "content": prompt
        })
        
        # Save to database with validation
        if database_available and user_id:
            from utils import update_conversation_title_if_needed
            session_id = st.session_state.get('imaginarium_session_id')
            if not session_id:
                print("ERROR: imaginarium_session_id is not set! Cannot save message.")
                st.warning("⚠️ Session error - message not saved. Please refresh the page.")
            else:
                db = get_db()
                if db:
                    try:
                        save_conversation_message(
                            db,
                            session_id=session_id,
                            interface_type='imaginarium',
                            role='user',
                            content=prompt,
                            user_id=user_id,
                            student_id=None
                        )
                        st.toast("✓ Message saved", icon="💾")
                        
                        # Auto-title on first message
                        conv_id = st.session_state.get('imaginarium_current_conversation_id')
                        if conv_id and len(st.session_state.imaginarium_messages) == 1:
                            update_conversation_title_if_needed(db, conv_id, 'imaginarium', prompt)
                    except Exception as e:
                        print(f"Error saving conversation: {str(e)}")
                        st.warning("⚠️ Unable to save message. Please check your connection.")
                    finally:
                        db.close()
        
        st.rerun()
    
    # Add scroll to top button
    add_scroll_to_top_button()


def show_contact_form():
    """Contact form with automatic email reply"""
    import requests
    from database import save_contact_submission, mark_contact_autoreply_sent, session_scope
    from auth import get_api_headers
    
    st.markdown("### 📬 Contact Us")
    st.markdown("*Have a question, feedback, or need assistance? We'd love to hear from you!*")
    
    st.markdown("---")
    
    # Get user info if logged in
    user_id = st.session_state.get('user_id')
    user_name = st.session_state.get('user_name', '')
    user_email = st.session_state.get('user_email', '')
    
    with st.form("contact_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        
        with col1:
            name = st.text_input("Your Name *", value=user_name, placeholder="Enter your name")
        
        with col2:
            email = st.text_input("Your Email *", value=user_email, placeholder="you@example.com")
        
        subject = st.text_input("Subject", placeholder="What is your message about?")
        
        message = st.text_area(
            "Your Message *", 
            placeholder="Please share your question, feedback, or how we can help...",
            height=200
        )
        
        submitted = st.form_submit_button("Send Message", use_container_width=True, type="primary")
        
        if submitted:
            if not name or not name.strip():
                st.error("Please enter your name.")
            elif not email or not email.strip():
                st.error("Please enter your email address.")
            elif '@' not in email or '.' not in email:
                st.error("Please enter a valid email address.")
            elif not message or not message.strip():
                st.error("Please enter your message.")
            else:
                try:
                    with session_scope() as db:
                        submission_id = save_contact_submission(
                            db,
                            name=name.strip(),
                            email=email.strip(),
                            subject=subject.strip() if subject else None,
                            message=message.strip(),
                            user_id=user_id
                        )
                        
                        if submission_id:
                            email_sent = False
                            try:
                                response = requests.post(
                                    'http://localhost:3001/api/email/send-contact-autoreply',
                                    json={
                                        'email': email.strip(),
                                        'userName': name.strip(),
                                        'subject': subject.strip() if subject else None
                                    },
                                    headers=get_api_headers(),
                                    timeout=10
                                )
                                
                                if response.status_code == 200:
                                    mark_contact_autoreply_sent(db, submission_id)
                                    email_sent = True
                            except Exception as e:
                                print(f"Error sending auto-reply email: {str(e)}")
                            
                            if email_sent:
                                st.success("Thank you for your message! We've sent you a confirmation email and will respond within 3 business days.")
                            else:
                                st.success("Thank you for your message! We've received it and will respond within 3 business days. (If you don't receive a confirmation email, please check your spam folder.)")
                        else:
                            st.error("Sorry, there was an issue submitting your message. Please try again.")
                            
                except Exception as e:
                    print(f"Error in contact form submission: {str(e)}")
                    st.error("Sorry, there was an issue submitting your message. Please try again.")
    
    st.markdown("---")
    
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.9em; padding: 20px;'>
        <p>Alternatively, you can email us directly at:<br>
        <strong>guide@auxpery.com.au</strong></p>
        <p style='margin-top: 15px; font-style: italic;'>
        As a small team, we appreciate your patience. We aim to respond to all messages within 3 business days.
        </p>
    </div>
    """, unsafe_allow_html=True)
